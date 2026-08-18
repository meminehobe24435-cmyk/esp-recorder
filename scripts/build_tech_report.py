# -*- coding: utf-8 -*-
"""
ESP Recorder 完整技术报告生成器
- 每个模块的完整源码 + 逐行拆解 + 设计思路
- 代码→功能对应关系 + 开发步骤
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===================== 完整项目源码 =====================

SRC = {}

SRC['app_main'] = r'''/* ===== main/app_main.c (183 行) =====
 * 主入口 — 按顺序初始化全部 13 个子系统
 * 核心: GPIO→NVS→SD→USB→Config→Recorder→UART→Router→WiFi→Web→Console
 */
#include "sd_card.h" "uart_driver.h" "usb_cdc.h" "usb_msc.h"
#include "wifi_config.h" "data_recorder.h" "data_router.h"
#include "config_manager.h" "web_server.h" "console_cmd.h"

#define TAG "main"
#define DEFAULT_BAUD_RATE 115200

/* ---- USB CDC 收到数据回调 (当前为空,保留扩展接口) ---- */
static void usb_cdc_rx_handler(const uint8_t *data, size_t len) {
    (void)data; (void)len;
}

/* ---- USB 插入 → 停止记录 → 卸载 FATFS → SD 交给 USB MSC ---- */
void tud_mount_cb(void) {
    ESP_LOGI(TAG, "USB mounted, releasing SD card from FATFS");
    recorder_stop();               // 刷新缓冲、关闭文件
    sd_card_unmount_fs();          // 卸载 VFS，保留块设备
}

/* ---- USB 拔出 → 重挂载 FATFS → 恢复记录 ---- */
void tud_umount_cb(void) {
    ESP_LOGI(TAG, "USB unmounted, reclaiming SD card for FATFS");
    if (sd_card_remount_fs() == ESP_OK) {
        const device_config_t *cfg = config_manager_get(config_manager_instance());
        if (cfg->recorder.auto_start) recorder_start();
    }
}

/* ---- 每个 UART 通道一个 RX 任务: 读→录→转 ---- */
static void uart_rx_task(void *arg) {
    uint8_t channel = (uint8_t)(uintptr_t)arg;
    uint8_t buf[256];
    while (1) {
        int len = user_uart_read_channel(channel, buf, sizeof(buf), 100);
        if (len > 0) {
            recorder_write_rx_channel(channel, buf, len);  // ①录
            data_router_forward(channel, buf, (size_t)len); // ②转
        }
    }
}

/* ========== 主入口 ========== */
void app_main(void) {
    // ① GPIO16 拉低 → 使能串口接收
    gpio_config_t uart_select = {
        .pin_bit_mask = 1 << GPIO_NUM_16, .mode = GPIO_MODE_OUTPUT, ...};
    gpio_config(&uart_select); gpio_set_level(GPIO_NUM_16, 0);

    // ② NVS 初始化 (WiFi 需要)
    esp_err_t ret = nvs_flash_init();
    if (ret == ESP_ERR_NVS_NO_FREE_PAGES || ret == ESP_ERR_NVS_NEW_VERSION_FOUND)
        { nvs_flash_erase(); nvs_flash_init(); }

    // ③ SD 卡挂载 (SPI → FATFS)
    sd_card_init();

    // ④ USB 复合设备 (CDC 虚拟串口 + MSC 读卡器)
    usb_cdc_init(usb_cdc_rx_handler); usb_msc_init();

    // ⑤ 加载 /sdcard/config.json，不存在则创建默认值
    config_manager_t *cfg_mgr = config_manager_instance();
    config_manager_load(cfg_mgr, "/sdcard/config.json");
    const device_config_t *cfg = config_manager_get(cfg_mgr);

    // ⑥ 记录引擎初始化 (环形缓冲 + IO/Sync 两个后台任务)
    recorder_init();

    // ⑦ 双路 UART 独立初始化 (一通道失败不阻塞另一通道)
    bool uart_ok[2] = {false, false};
    for (uint8_t ch = 1; ch <= 2; ch++)
        uart_ok[ch-1] = (user_uart_init_channel(ch, cfg->uart.baudrate) == ESP_OK);
    user_uart_set_tx_callback(recorder_write_tx_channel);
    data_router_init();  // ⑧ 跨通道路由

    // ⑨⑩ WiFi AP 热点
    wifi_config_init();
    wifi_apply_ap(cfg->wifi.ap_ssid, cfg->wifi.ap_pass, cfg->wifi.ap_ip);

    // ⑪ HTTP Web 服务器 (端口 80, 14 个 API)
    web_server_start();

    // ⑫ 可选: 连接外部 WiFi (STA)
    if (cfg->wifi.enable_sta && cfg->wifi.sta_ssid[0])
        wifi_connect_sta(cfg->wifi.sta_ssid, cfg->wifi.sta_pass, cfg->wifi.sta_authmode);

    // ⑫ 可选: 上电自动记录
    if (cfg->recorder.auto_start) recorder_start();

    // ⑬ 创建两路 UART RX 任务
    for (uint8_t ch = 1; ch <= 2; ch++)
        if (uart_ok[ch-1]) xTaskCreate(uart_rx_task, "uart_rx", 4096, (void*)ch, 5, NULL);

    console_start();  // 调试控制台
    ESP_LOGI(TAG, "esp_recorder init complete");
    while (1) vTaskDelay(pdMS_TO_TICKS(1000));
}'''

SRC['uart_h'] = r'''/* ===== main/uart_driver.h (47 行) =====
 * 两路独立用户串口定义
 * CH1=UART0(GPIO43/44)  CH2=UART2(GPIO17/18)
 * 调试控制台使用 UART1，不占用这两个通道 */
#ifndef UART_DRIVER_H
#define UART_DRIVER_H
#include <stdbool.h> <stdint.h> <stddef.h>
#include "esp_err.h" "driver/uart.h"

#define USER_UART_CHANNEL_COUNT 2
#define USER_UART_CHANNEL_1     1
#define USER_UART_CHANNEL_2     2

#define USER_UART1_NUM       UART_NUM_0
#define USER_UART1_TX_PIN    43
#define USER_UART1_RX_PIN    44

#define USER_UART2_NUM       UART_NUM_2
#define USER_UART2_TX_PIN    17
#define USER_UART2_RX_PIN    18

#define USER_UART_BUF_SIZE   2048

esp_err_t user_uart_init(int baud_rate);
esp_err_t user_uart_init_channel(uint8_t channel, int baud_rate);
esp_err_t user_uart_set_baudrate(int baud_rate);
esp_err_t user_uart_set_channel_baudrate(uint8_t channel, int baud_rate);

int user_uart_write(const uint8_t *data, size_t len);
int user_uart_read(uint8_t *buf, size_t len, uint32_t timeout_ms);
int user_uart_write_channel(uint8_t channel, const uint8_t *data, size_t len);
int user_uart_read_channel(uint8_t channel, uint8_t *buf, size_t len, uint32_t t_ms);

bool user_uart_channel_is_ready(uint8_t channel);
uint64_t user_uart_get_rx_bytes(uint8_t channel);
uint64_t user_uart_get_tx_bytes(uint8_t channel);

typedef esp_err_t (*user_uart_tx_cb_t)(uint8_t channel, const uint8_t *data, size_t len);
void user_uart_set_tx_callback(user_uart_tx_cb_t cb);
#endif'''

SRC['uart_c'] = r'''/* ===== main/uart_driver.c (221 行) =====
 * 每个函数的作用和实现逻辑 */
#include "uart_driver.h"
#include "freertos/FreeRTOS.h" "driver/uart.h" "esp_log.h"
#define TAG "uart_driver"

/* ---- 全局状态 (portMUX_TYPE 自旋锁保证多任务线程安全) ---- */
static user_uart_tx_cb_t s_tx_cb = NULL;
static portMUX_TYPE s_state_lock = portMUX_INITIALIZER_UNLOCKED;
static bool     s_channel_ready[2];
static uint64_t s_rx_bytes[2], s_tx_bytes[2];

/* ---- 硬件映射表 ---- */
typedef struct { uart_port_t port; int tx_pin, rx_pin; } user_uart_hw_t;
static const user_uart_hw_t s_uart_hw[2] = {
    {USER_UART1_NUM, USER_UART1_TX_PIN, USER_UART1_RX_PIN},  // CH1: UART0, GPIO43/44
    {USER_UART2_NUM, USER_UART2_TX_PIN, USER_UART2_RX_PIN},  // CH2: UART2, GPIO17/18
};

/* ---- 单通道初始化 (核心函数) ---- */
esp_err_t user_uart_init_channel(uint8_t channel, int baud_rate) {
    const user_uart_hw_t *hw = &s_uart_hw[channel - 1];

    if (!uart_is_driver_installed(hw->port)) {
        // 首次: 配置参数 → 设置引脚 → 安装驱动
        uart_config_t cfg = { .baud_rate = baud_rate, .data_bits = UART_DATA_8_BITS,
            .parity = UART_PARITY_DISABLE, .stop_bits = UART_STOP_BITS_1,
            .flow_ctrl = UART_HW_FLOWCTRL_DISABLE, .source_clk = UART_SCLK_DEFAULT };
        uart_param_config(hw->port, &cfg);
        uart_set_pin(hw->port, hw->tx_pin, hw->rx_pin, UART_PIN_NO_CHANGE, UART_PIN_NO_CHANGE);
        uart_driver_install(hw->port, USER_UART_BUF_SIZE, USER_UART_BUF_SIZE, 0, NULL, 0);
    } else {
        // 已有驱动: 只改波特率
        uart_set_baudrate(hw->port, baud_rate);
    }
    s_channel_ready[channel - 1] = true;  // 标记就绪
    return ESP_OK;
}

/* ---- 写入 (发送) 一个通道 ---- */
int user_uart_write_channel(uint8_t channel, const uint8_t *data, size_t len) {
    const user_uart_hw_t *hw = &s_uart_hw[channel - 1];
    int sent = uart_write_bytes(hw->port, data, len);
    if (sent > 0) {
        portENTER_CRITICAL(&s_state_lock);
        s_tx_bytes[channel - 1] += sent;  // 累计 TX 字节数
        portEXIT_CRITICAL(&s_state_lock);
        if (s_tx_cb) s_tx_cb(channel, data, sent);  // 回调 (写入实时流)
    }
    return sent;
}

/* ---- 读取 (接收) 一个通道 ---- */
int user_uart_read_channel(uint8_t channel, uint8_t *buf, size_t len, uint32_t timeout_ms) {
    const user_uart_hw_t *hw = &s_uart_hw[channel - 1];
    int rx = uart_read_bytes(hw->port, buf, len, pdMS_TO_TICKS(timeout_ms));
    if (rx > 0) {
        portENTER_CRITICAL(&s_state_lock);
        s_rx_bytes[channel - 1] += rx;  // 累计 RX 字节数
        portEXIT_CRITICAL(&s_state_lock);
    }
    return rx;
}

/* ---- 状态查询 (线程安全) ---- */
bool user_uart_channel_is_ready(uint8_t ch) { return s_channel_ready[ch-1]; }
uint64_t user_uart_get_rx_bytes(uint8_t ch) { return s_rx_bytes[ch-1]; }
uint64_t user_uart_get_tx_bytes(uint8_t ch) { return s_tx_bytes[ch-1]; }
void user_uart_set_tx_callback(user_uart_tx_cb_t cb) { s_tx_cb = cb; }'''

SRC['sd_card'] = r'''/* ===== main/sd_card.c (171 行) =====
 * SD 卡通过 SPI 总线访问，挂载 FATFS 到 /sdcard
 * 支持安全卸载(留给USB MSC)和安全重挂载(USB拔出后) */
#include "sd_card.h"
#include "driver/sdspi_host.h" "driver/spi_common.h"
#include "sdmmc_cmd.h" "esp_vfs_fat.h" "diskio_sdmmc.h" "diskio_impl.h" "ff.h"
#define TAG "sd_card"

#define SD_MOUNT_POINT      "/sdcard"
#define SD_HOST             SPI2_HOST
#define SD_PIN_CS           10
#define SD_PIN_MOSI         38
#define SD_PIN_MISO         40
#define SD_PIN_SCK          39
#define SD_MAX_TRANSFER_SZ  4000

static sd_card_state_t s_state = {0};
static sdmmc_card_t *s_card = NULL;
static bool s_spi_initialized = false;

/* ---- 初始化: SPI 总线 → SDSPI 主机 → 挂载 FATFS ---- */
esp_err_t sd_card_init(void) {
    if (s_state.mounted) return ESP_OK;
    // 1. 初始化 SPI 总线
    if (!s_spi_initialized) {
        spi_bus_config_t bus = { .mosi_io_num = SD_PIN_MOSI, .miso_io_num = SD_PIN_MISO,
            .sclk_io_num = SD_PIN_SCK, .max_transfer_sz = SD_MAX_TRANSFER_SZ };
        spi_bus_initialize(SD_HOST, &bus, SDSPI_DEFAULT_DMA);
        s_spi_initialized = true;
    }
    // 2. SDSPI 主机 (20MHz)
    sdmmc_host_t host = SDSPI_HOST_DEFAULT();
    host.slot = SD_HOST; host.max_freq_khz = 20000;
    // 3. 设备配置 (CS 引脚)
    sdspi_device_config_t slot = SDSPI_DEVICE_CONFIG_DEFAULT();
    slot.gpio_cs = SD_PIN_CS; slot.host_id = SD_HOST;
    // 4. 挂载 FATFS (不自动格式化)
    esp_vfs_fat_sdmmc_mount_config_t mnt = {
        .format_if_mount_failed = false, .max_files = 5, .allocation_unit_size = 16*1024 };
    esp_err_t ret = esp_vfs_fat_sdspi_mount(SD_MOUNT_POINT, &host, &slot, &mnt, &s_card);
    if (ret == ESP_OK) { s_state.mounted = true;
        ESP_LOGI(TAG, "SD mounted, size=%llu MB",
            (uint64_t)s_card->csd.capacity * s_card->csd.sector_size / (1024*1024)); }
    return ret;
}

/* ---- 只卸载 FATFS，保留 SPI 和 sdmmc_card_t (给 USB MSC 用) ---- */
esp_err_t sd_card_unmount_fs(void) {
    if (!s_state.mounted) return ESP_OK;
    BYTE pdrv = ff_diskio_get_pdrv_card(s_card);
    if (pdrv != FF_DRV_NOT_USED) {
        char drv[3] = {(char)('0'+pdrv), ':', 0};
        f_mount(0, drv, 0);           // 卸载 FatFs 逻辑卷
        ff_diskio_unregister(pdrv);   // 注销 diskio
    }
    esp_vfs_fat_unregister_path(SD_MOUNT_POINT);  // 注销 VFS
    s_state.mounted = false;
    return ESP_OK;
}

/* ---- USB 拔出后重新挂载 FATFS ---- */
esp_err_t sd_card_remount_fs(void) {
    sd_card_deinit();    // 完全释放 (SPI + card)
    return sd_card_init(); // 重新完整初始化
}

sdmmc_card_t *sd_card_get_handle(void) { return s_card; }
const sd_card_state_t *sd_card_get_state(void) { return &s_state; }'''

SRC['usb_cdc'] = r'''/* ===== main/usb_cdc.c (150 行) =====
 * TinyUSB CDC-ACM 虚拟串口 + MSC 复合设备描述符
 * USB 插入时枚举为: CDC(虚拟串口) + MSC(SD读卡器) */
#include "usb_cdc.h"
#include "tinyusb.h" "tusb_cdc_acm.h" "esp_log.h" "driver/gpio.h"
#define TAG "usb_cdc"
#define USB_CDC_TX_BUF_SIZE 512
#define USB_CDC_RX_BUF_SIZE 512

enum { ITF_NUM_CDC=0, ITF_NUM_CDC_DATA, ITF_NUM_MSC, ITF_NUM_TOTAL };
enum { EDPT_CDC_NOTIF=0x81, EDPT_CDC_OUT=0x02, EDPT_CDC_IN=0x82,
       EDPT_MSC_OUT=0x03, EDPT_MSC_IN=0x83 };

/* ---- 配置描述符: CDC控制接口 + CDC数据接口 + MSC接口 ---- */
static uint8_t const desc_configuration[] = {
    TUD_CONFIG_DESCRIPTOR(1, ITF_NUM_TOTAL, 0, TUSB_DESC_TOTAL_LEN,
                          TUSB_DESC_CONFIG_ATT_REMOTE_WAKEUP, 100),
    TUD_CDC_DESCRIPTOR(ITF_NUM_CDC, 4, EDPT_CDC_NOTIF, 8,
                       EDPT_CDC_OUT, EDPT_CDC_IN, 64),
    TUD_MSC_DESCRIPTOR(ITF_NUM_MSC, 5, EDPT_MSC_OUT, EDPT_MSC_IN,
                       TUD_OPT_HIGH_SPEED ? 512 : 64),
};

/* ---- 设备描述符 ---- */
static tusb_desc_device_t descriptor_config = {
    .bLength=sizeof(descriptor_config), .bDescriptorType=TUSB_DESC_DEVICE,
    .bcdUSB=0x0200, .bDeviceClass=TUSB_CLASS_MISC,
    .idVendor=0x303A, .idProduct=0x4002, .bcdDevice=0x100,
    .iManufacturer=0x01, .iProduct=0x02, .iSerialNumber=0x03, .bNumConfigurations=1
};

static const char *s_string_desc[] = {
    (const char[]){0x09,0x04},  // Langid
    "Espressif", "ESP Recorder", "123456", "ESP Recorder CDC", "ESP Recorder MSC"
};

/* ---- CDC 接收回调 ---- */
static void usb_cdc_rx_callback(int itf, cdcacm_event_t *event) {
    uint8_t buf[USB_CDC_RX_BUF_SIZE]; size_t rx_size = 0;
    while (tinyusb_cdcacm_read(TINYUSB_CDC_ACM_0, buf, sizeof(buf), &rx_size) == ESP_OK
           && rx_size > 0) {
        if (s_rx_cb) s_rx_cb(buf, rx_size);  // 回调注册的处理函数
    }
}

/* ---- 初始化 TinyUSB (CDC + MSC 复合设备) ---- */
esp_err_t usb_cdc_init(usb_cdc_rx_cb_t rx_cb) {
    s_rx_cb = rx_cb;
    tinyusb_config_t tusb_cfg = { .device_descriptor = &descriptor_config,
        .string_descriptor = s_string_desc, .string_descriptor_count = 6,
        .external_phy = false, .configuration_descriptor = desc_configuration,
        .self_powered = true, .vbus_monitor_io = GPIO_NUM_48 };
    tinyusb_driver_install(&tusb_cfg);
    tinyusb_config_cdcacm_t cdc_cfg = { .usb_dev = TINYUSB_USBDEV_0,
        .cdc_port = TINYUSB_CDC_ACM_0, .rx_unread_buf_sz = USB_CDC_RX_BUF_SIZE,
        .callback_rx = &usb_cdc_rx_callback };
    tusb_cdc_acm_init(&cdc_cfg);
    return ESP_OK;
}

int usb_cdc_write(const uint8_t *data, size_t len) {
    size_t q = tinyusb_cdcacm_write_queue(TINYUSB_CDC_ACM_0, data, len);
    if (q > 0) tinyusb_cdcacm_write_flush(TINYUSB_CDC_ACM_0, 0);
    return (int)q;
}'''

SRC['usb_msc'] = r'''/* ===== main/usb_msc.c (189 行) =====
 * USB MSC (大容量存储) — 把 SD 卡暴露为电脑 U 盘
 * 实现 SCSI INQUIRY / READ10 / WRITE10 / START_STOP 回调 */
#include "usb_msc.h"
#include "tusb.h" "sd_card.h" "sdmmc_cmd.h"
#define TAG "usb_msc"
#define MSC_IO_RETRY_MAX 3     // 读写失败最多重试 3 次
#define MSC_IO_RETRY_DELAY_MS 10

static sdmmc_card_t *msc_get_card(void) { return sd_card_get_handle(); }

/* ---- SCSI INQUIRY: 告诉主机我是 "Espressif ESP Recorder MSC v1.0" ---- */
void tud_msc_inquiry_cb(uint8_t lun, uint8_t vid[8], uint8_t pid[16], uint8_t rev[4]) {
    memcpy(vid, "Espressif", 9); memcpy(pid, "ESP Recorder MSC", 16);
    memcpy(rev, "1.0", 3);
}

/* ---- SCSI TEST UNIT READY: SD 卡在不在? ---- */
bool tud_msc_test_unit_ready_cb(uint8_t lun) { return msc_get_card() != NULL; }

/* ---- SCSI READ CAPACITY: 告诉主机 SD 卡多大 ---- */
void tud_msc_capacity_cb(uint8_t lun, uint32_t *blocks, uint16_t *bs) {
    sdmmc_card_t *card = msc_get_card();
    *blocks = card->csd.capacity; *bs = (uint16_t)card->csd.sector_size;
}

/* ---- SCSI READ10: 主机读扇区 (带 3 次重试) ---- */
int32_t tud_msc_read10_cb(uint8_t lun, uint32_t lba, uint32_t offset,
                          void *buf, uint32_t bufsize) {
    sdmmc_card_t *card = msc_get_card();
    for (int i = 0; i < MSC_IO_RETRY_MAX; i++) {
        if (sdmmc_read_sectors(card, buf, lba, bufsize / card->csd.sector_size) == ESP_OK)
            return (int32_t)bufsize;
        vTaskDelay(pdMS_TO_TICKS(MSC_IO_RETRY_DELAY_MS));
    }
    return -1;
}

/* ---- SCSI WRITE10: 主机写扇区 (带 3 次重试) ---- */
int32_t tud_msc_write10_cb(uint8_t lun, uint32_t lba, uint32_t offset,
                           uint8_t *buf, uint32_t bufsize) {
    sdmmc_card_t *card = msc_get_card();
    for (int i = 0; i < MSC_IO_RETRY_MAX; i++) {
        if (sdmmc_write_sectors(card, buf, lba, bufsize / card->csd.sector_size) == ESP_OK)
            return (int32_t)bufsize;
        vTaskDelay(pdMS_TO_TICKS(MSC_IO_RETRY_DELAY_MS));
    }
    return -1;
}

esp_err_t usb_msc_init(void) {
    ESP_LOGI(TAG, "USB MSC callbacks registered");
    return ESP_OK;
}'''

SRC['wifi_config'] = r'''/* ===== main/wifi_config.c (253 行) =====
 * WiFi AP+STA 混合模式管理
 * AP: 默认 SSID=ESP_Recorder_XXXX, 密码=12345678, IP=192.168.4.1
 * STA: 支持 WPA2/WPA3/开放, 断线自动重连 3 次间隔 5 秒 */
#include "wifi_config.h"
#include "freertos/FreeRTOS.h" "freertos/event_groups.h"
#include "esp_wifi.h" "esp_event.h" "esp_netif.h" "esp_mac.h" "lwip/inet.h"
#define TAG "wifi_config"
#define STA_MAX_RETRY 3
#define STA_RETRY_INTERVAL_MS 5000

static EventGroupHandle_t s_wifi_event_group;
static wifi_state_t s_state = {0};
static int s_retry_count = 0;

/* ---- 生成默认 SSID: ESP_Recorder_XXYY (MAC 后两位) ---- */
static char *wifi_ap_default_ssid(char *out, size_t sz) {
    uint8_t mac[6]; esp_read_mac(mac, ESP_MAC_WIFI_STA);
    snprintf(out, sz, "ESP_Recorder_%02X%02X", mac[4], mac[5]); return out;
}

/* ---- WiFi 事件处理器 ---- */
static void wifi_event_handler(void *arg, esp_event_base_t base, int32_t id, void *data) {
    if (base == WIFI_EVENT && id == WIFI_EVENT_STA_DISCONNECTED) {
        s_state.sta_connected = false;
        if (s_retry_count < STA_MAX_RETRY) {  // 自动重连 (最多 3 次)
            s_retry_count++;
            vTaskDelay(pdMS_TO_TICKS(STA_RETRY_INTERVAL_MS));
            esp_wifi_connect();
        }
    } else if (base == IP_EVENT && id == IP_EVENT_STA_GOT_IP) {
        s_retry_count = 0; s_state.sta_connected = true;  // 重连成功, 计数清零
    } else if (base == WIFI_EVENT && id == WIFI_EVENT_AP_START) {
        s_state.ap_started = true;
    }
}

/* ---- 初始化 WiFi 协议栈 ---- */
esp_err_t wifi_config_init(void) {
    s_wifi_event_group = xEventGroupCreate();
    esp_netif_init(); esp_event_loop_create_default();
    esp_netif_create_default_wifi_ap();    // AP 接口
    esp_netif_create_default_wifi_sta();   // STA 接口
    wifi_init_config_t cfg = WIFI_INIT_CONFIG_DEFAULT();
    esp_wifi_init(&cfg);
    esp_event_handler_instance_register(WIFI_EVENT, ESP_EVENT_ANY_ID, &wifi_event_handler, NULL, NULL);
    esp_event_handler_instance_register(IP_EVENT, IP_EVENT_STA_GOT_IP, &wifi_event_handler, NULL, NULL);
    return ESP_OK;
}

/* ---- 启动/更新 AP 热点 ---- */
esp_err_t wifi_apply_ap(const char *ssid, const char *pass, const char *ip) {
    // 设置静态 IP
    esp_netif_ip_info_t info = { .ip.addr = ipaddr_addr(ip ? ip : "192.168.4.1"),
        .gw.addr = ipaddr_addr("192.168.4.1"), .netmask.addr = ipaddr_addr("255.255.255.0") };
    esp_netif_dhcps_stop(s_ap_netif); esp_netif_set_ip_info(s_ap_netif, &info);
    esp_netif_dhcps_start(s_ap_netif);
    // 构造 AP 配置
    wifi_config_t ap = { .ap = { .channel = 1, .max_connection = 4,
        .authmode = (pass && pass[0]) ? WIFI_AUTH_WPA2_PSK : WIFI_AUTH_OPEN } };
    strncpy((char*)ap.ap.ssid, (ssid&&ssid[0])?ssid:wifi_ap_default_ssid(...), 31);
    if (pass) strncpy((char*)ap.ap.password, pass, 63);
    esp_wifi_set_mode(WIFI_MODE_APSTA); esp_wifi_set_config(WIFI_IF_AP, &ap);
    esp_wifi_start();
    return ESP_OK;
}

/* ---- 连接外部 WiFi (STA) ---- */
esp_err_t wifi_connect_sta(const char *ssid, const char *pass, int authmode) {
    wifi_config_t sta = {0};
    strncpy((char*)sta.sta.ssid, ssid, 31);
    strncpy((char*)sta.sta.password, pass ? pass : "", 63);
    sta.sta.threshold.authmode = (wifi_auth_mode_t)authmode;
    s_retry_count = 0;
    esp_wifi_disconnect(); esp_wifi_set_config(WIFI_IF_STA, &sta);
    if (!s_state.ap_started) esp_wifi_start();
    return ESP_OK;
}'''

SRC['data_recorder'] = r'''/* ===== main/data_recorder.c (614 行) — 最核心模块 =====
 * 双通道独立数据记录引擎
 *
 * 架构: RX数据 → live_push(实时流256条) + ringbuf_write(落盘16KB缓冲)
 *                     ↓
 *         rec_io 任务: 阈值≥2KB 或 超时1秒 → fwrite(SD卡)
 *         rec_sync 任务: 每秒 fsync() 强制同步, 连续3次失败→关闭文件
 *
 * 文件命名: /sdcard/REC_CH1_20260727_193428.bin (同名自动加_1后缀)
 * 实时流格式: 2026-07-27T19:34:28.123,CH1,RX,01 03 00 00 00 0A C5 CD */
#include "data_recorder.h"
#include "freertos/FreeRTOS.h" "freertos/semphr.h" "sd_card.h"
#define TAG "recorder"

#define SAVE_RINGBUF_SIZE   16384   // 每通道 16KB 环形缓冲
#define SAVE_THRESHOLD_BYTES 2048   // 满 2KB 立即刷盘
#define SAVE_TIMEOUT_MS     1000    // 超时 1 秒也刷盘
#define SAVE_TASK_STACK      4096
#define SYNC_TASK_STACK      3072
#define LIVE_CHUNK_BYTES       48   // 实时流每行 HEX 块大小

/* ---- 环形缓冲结构 ---- */
typedef struct { uint8_t *buf; size_t size, head, tail; } ringbuf_t;

/* ---- 每通道状态 ---- */
typedef struct { FILE *file; char filename[96]; ringbuf_t save_rb;
    uint64_t rx_total; uint32_t drop_count, fsync_fail_count; bool failed;
} recorder_channel_state_t;

static bool s_running = false;
static SemaphoreHandle_t s_mutex, s_data_sem, s_live_mutex;
static recorder_channel_state_t s_channels[2];
static recorder_live_item_t s_live[256];  // 实时流环形缓冲
static uint32_t s_live_seq = 0;

/* ---- 环形缓冲读写 (标准实现) ---- */
static size_t ringbuf_write(ringbuf_t *rb, const uint8_t *data, size_t len) {
    size_t free = (rb->tail + rb->size - rb->head - 1) % rb->size;
    if (len > free) len = free;
    for (size_t i = 0; i < len; i++)
        { rb->buf[rb->head] = data[i]; rb->head = (rb->head + 1) % rb->size; }
    return len;
}

/* ---- 落盘任务: 阈值/超时双条件触发 ---- */
static void recorder_io_task(void *arg) {
    while (1) {
        // 等待: 有新数据 或 超时 1 秒
        bool timeout = (xSemaphoreTake(s_data_sem, pdMS_TO_TICKS(SAVE_TIMEOUT_MS)) == pdFALSE);
        xSemaphoreTake(s_mutex, portMAX_DELAY);
        bool flush = timeout;
        if (!flush)
            for (int ch = 0; ch < 2; ch++)
                if (ringbuf_count(&s_channels[ch].save_rb) >= SAVE_THRESHOLD_BYTES)
                    { flush = true; break; }
        if (flush) recorder_flush_all_locked();  // fwrite → SD 卡
        xSemaphoreGive(s_mutex);
    }
}

/* ---- Sync 任务: 每秒 fsync() ---- */
static void recorder_sync_task(void *arg) {
    while (1) {
        vTaskDelay(pdMS_TO_TICKS(1000));
        xSemaphoreTake(s_mutex, portMAX_DELAY);
        for (int ch = 0; ch < 2; ch++) {
            if (!s_channels[ch].file) continue;
            int rc = fsync(fileno(s_channels[ch].file));
            if (rc != 0 && ++s_channels[ch].fsync_fail_count >= 3) {
                s_channels[ch].failed = true;   // 连续 3 次失败 → 关闭文件
                recorder_close_channel_locked(ch + 1);
            }
        }
        xSemaphoreGive(s_mutex);
    }
}

/* ---- 写入 RX 数据: 实时流 + 落盘 ---- */
esp_err_t recorder_write_rx_channel(uint8_t channel, const uint8_t *data, size_t len) {
    recorder_live_push(channel, 0, data, len);  // ① 实时流
    recorder_write_raw(channel, data, len);      // ② 落盘环形缓冲
    return ESP_OK;
}

/* ---- 开始记录: 清空统计 → 预计算文件名 → 标记运行 ---- */
esp_err_t recorder_start(void) {
    for (int ch = 0; ch < 2; ch++) {
        ringbuf_reset(&s_channels[ch].save_rb);
        s_channels[ch].rx_total = s_channels[ch].drop_count = 0;
        s_channels[ch].fsync_fail_count = 0; s_channels[ch].failed = false;
        recorder_compute_filename_locked(ch + 1);  // 生成 REC_CH1_时间戳.bin
    }
    s_running = true;
    return ESP_OK;
}

/* ---- 停止记录: flush 所有缓冲 → 关闭文件 ---- */
esp_err_t recorder_stop(void) {
    s_running = false;
    recorder_flush_all_locked();
    for (int ch = 0; ch < 2; ch++) recorder_close_channel_locked(ch + 1);
    return ESP_OK;
}'''

SRC['data_router'] = r'''/* ===== main/data_router.c (79 行) =====
 * 跨通道数据路由: CH1 RX → CH2 TX, CH2 RX → CH1 TX
 * 双向转发 + 丢包统计, 线程安全 */
#include "data_router.h" "uart_driver.h"
#define TAG "data_router"

static portMUX_TYPE s_lock = portMUX_INITIALIZER_UNLOCKED;
static data_router_stats_t s_stats;

esp_err_t data_router_init(void) {
    memset(&s_stats, 0, sizeof(s_stats)); s_stats.enabled = true;
    return ESP_OK;
}

int data_router_forward(uint8_t src, const uint8_t *data, size_t len) {
    // 确定目标通道: CH1→CH2, CH2→CH1
    uint8_t dst = (src == USER_UART_CHANNEL_1) ? USER_UART_CHANNEL_2 : USER_UART_CHANNEL_1;
    if (!s_stats.enabled) return 0;

    int sent = user_uart_write_channel(dst, data, len);
    // 更新统计
    portENTER_CRITICAL(&s_lock);
    if (src == USER_UART_CHANNEL_1) { s_stats.ch1_to_ch2_bytes += sent;
        s_stats.ch1_to_ch2_drop_bytes += (len - (sent>0?sent:0)); }
    else { s_stats.ch2_to_ch1_bytes += sent;
        s_stats.ch2_to_ch1_drop_bytes += (len - (sent>0?sent:0)); }
    portEXIT_CRITICAL(&s_lock);
    return sent;
}

void data_router_get_stats(data_router_stats_t *stats) {
    portENTER_CRITICAL(&s_lock); *stats = s_stats; portEXIT_CRITICAL(&s_lock);
}'''

SRC['web_server'] = r'''/* ===== main/web_server.c (649 行) =====
 * HTTP REST API 服务器 — 14 个端点
 * 路径穿越防护 + IPv4 校验 + cJSON 安全拼装 JSON */
#include "web_server.h"
#include "esp_http_server.h" "cJSON.h"
#include "wifi_config.h" "uart_driver.h" "data_recorder.h" "data_router.h"
#include "config_manager.h" "sd_card.h"
#define TAG "web_server"
#define WEB_PORT 80

/* ==== 14 个 API 路由表 ==== */
static const httpd_uri_t uris[] = {
    {"/",                     HTTP_GET,  index_handler},            //  1. 首页
    {"/api/status",           HTTP_GET,  status_handler},           //  2. 状态
    {"/api/wifi",             HTTP_POST, wifi_set_handler},         //  3. STA配网
    {"/api/ap",               HTTP_POST, ap_set_handler},           //  4. AP配置
    {"/api/files",            HTTP_GET,  files_list_handler},       //  5. 文件列表
    {"/api/file",             HTTP_GET,  file_download_handler},    //  6. 下载
    {"/api/file/delete",      HTTP_POST, file_delete_handler},      //  7. 删除
    {"/api/recorder/start",   HTTP_POST, recorder_start_handler},   //  8. 开始记录
    {"/api/recorder/stop",    HTTP_POST, recorder_stop_handler},    //  9. 停止记录
    {"/api/stream",           HTTP_GET,  stream_fetch_handler},     // 10. 实时流
    {"/api/stream/clear",     HTTP_POST, stream_clear_handler},     // 11. 清空流
    {"/api/send",             HTTP_POST, send_tx_handler},          // 12. HEX发送
    {"/api/reboot",           HTTP_POST, reboot_handler},           // 13. 重启
    {"/api/config/stream_gap",HTTP_POST, stream_gap_set_handler},   // 14. 设间隔
};

/* ==== 首页: 返回内嵌的 index.html ==== */
static esp_err_t index_handler(httpd_req_t *req) {
    extern const uint8_t _binary_index_html_start[], _binary_index_html_end[];
    size_t len = _binary_index_html_end - _binary_index_html_start;
    httpd_resp_set_type(req, "text/html; charset=utf-8");
    return httpd_resp_send(req, (const char*)_binary_index_html_start, len);
}

/* ==== 状态 API: 返回完整 JSON (25+ 字段) ==== */
static esp_err_t status_handler(httpd_req_t *req) {
    cJSON *root = cJSON_CreateObject();
    cJSON_AddBoolToObject(root, "ap", ws->ap_started);
    cJSON_AddBoolToObject(root, "sta_connected", ws->sta_connected);
    cJSON_AddBoolToObject(root, "sd", sd->mounted);
    cJSON_AddBoolToObject(root, "recording", recorder_is_running());
    cJSON_AddNumberToObject(root, "rec_drop_count", recorder_get_drop_count());
    cJSON_AddNumberToObject(root, "rec_total_bytes", (double)recorder_get_total_bytes());
    cJSON_AddBoolToObject(root, "route_enabled", route_stats.enabled);
    // ... 共 25+ 字段
    char *body = cJSON_PrintUnformatted(root);
    httpd_resp_set_type(req, "application/json; charset=utf-8");
    return httpd_resp_send(req, body, strlen(body));
}

/* ==== 实时流 API: 增量拉取 since_seq 之后的数据 ==== */
static esp_err_t stream_fetch_handler(httpd_req_t *req) {
    uint32_t since = 0; get_query_value(req, "since", &since);
    recorder_live_item_t items[32]; uint32_t max_seq;
    size_t got = recorder_live_fetch(since, items, 32, &max_seq);
    // 返回 JSON: {"items":[{"seq":1,"ts":...,"line":"...RX,01 03..."}],"max_seq":5}
}

/* ==== HEX 发送 API: POST {"channel":1,"hex":"01 03 00 0A"} ==== */
static esp_err_t send_tx_handler(httpd_req_t *req) {
    // 解析 JSON body → 解码 HEX 字符串 → uart_write → 返回 {"ok":true,"sent":6}
}

/* ==== 启动 HTTP 服务器 ==== */
esp_err_t web_server_start(void) {
    httpd_config_t cfg = HTTPD_DEFAULT_CONFIG();
    cfg.server_port = WEB_PORT; cfg.max_uri_handlers = 16; cfg.stack_size = 16384;
    httpd_start(&s_server, &cfg);
    for (int i = 0; i < 14; i++) httpd_register_uri_handler(s_server, &uris[i]);
    return ESP_OK;
}'''

SRC['config_manager'] = r'''/* ===== main/config_manager.c (278 行) =====
 * JSON 配置系统 — 从 /sdcard/config.json 加载, 缺失字段用默认值
 * authmode 合法性校验, stream_gap_ms clamp 到 [1,1000], 写盘失败自动回滚 */
#include "config_manager.h" "cJSON.h" "esp_wifi.h"
#define TAG "config" #define CONFIG_MAX_SIZE 4096

/* authmode 合法性校验: 只允许 OPEN/WPA/WPA2/WPA3/WPA2_WPA3 */
static int clamp_authmode(int v) {
    switch (v) { case 0: case 2: case 3: case 6: case 7: return v; default: return 0; }
}

/* 加载配置: 默认值 → fopen → cJSON_Parse → 逐字段覆盖 */
esp_err_t config_manager_load(config_manager_t* mgr, const char* path) {
    config_manager_reset_defaults(mgr);   // 先填默认值
    FILE* f = fopen(path, "r"); if (!f) return ESP_ERR_NOT_FOUND;
    char* buf = malloc(size+1); fread(buf, 1, size, f); fclose(f);
    cJSON* root = cJSON_Parse(buf);       // 解析 JSON
    // UART 段
    cJSON* uart = cJSON_GetObjectItem(root, "uart");
    if (uart) { mgr->cfg.uart.baudrate = get_json_int(uart, "baudrate", 115200); ... }
    // WiFi 段
    cJSON* wifi = cJSON_GetObjectItem(root, "wifi");
    if (wifi) { mgr->cfg.wifi.sta_authmode = clamp_authmode(get_json_int(wifi, "sta_authmode", 0)); ... }
    // Recorder 段
    cJSON* rec = cJSON_GetObjectItem(root, "recorder");
    if (rec) { mgr->cfg.recorder.auto_start = get_json_int(rec, "auto_start", 1) != 0;
        mgr->cfg.recorder.stream_gap_ms = clamp_gap(get_json_int(rec, "stream_gap_ms", 100)); }
    mgr->loaded = true; return ESP_OK;
}

/* 保存配置: cJSON 拼装 → fopen("w") → fputs */
esp_err_t config_manager_save(const config_manager_t* mgr, const char* path) {
    cJSON* root = cJSON_CreateObject();
    cJSON* uart = cJSON_CreateObject();
    cJSON_AddNumberToObject(uart, "baudrate", mgr->cfg.uart.baudrate);
    // ... WiFi段, Recorder段
    char* json = cJSON_Print(root);
    FILE* f = fopen(path, "w"); fputs(json, f); fclose(f);
    cJSON_free(json); cJSON_Delete(root);
    return ESP_OK;
}'''

SRC['console'] = r'''/* ===== main/console_cmd.c (91 行) =====
 * 串口调试控制台 — UART1 (TX=GPIO1, RX=GPIO2, 115200 baud)
 * 基于 ESP-IDF Console REPL 组件, 注册 sd_test 命令 */
#include "console_cmd.h" "sd_speed_test.h"
#include "esp_console.h" "linenoise/linenoise.h"
#define TAG "console"

/* sd_test 命令处理 */
static int cmd_sd_test(int argc, char **argv) {
    if (argc == 1) return sd_speed_test_run_all() == ESP_OK ? 0 : 1;           // 无参: 标准四档
    if (argc == 3) { sd_speed_result_t r;                                       // 有参: 自定义
        sd_speed_test_once(strtol(argv[1],NULL,10), strtol(argv[2],NULL,10), &r);
        printf("block=%u total=%u write=%.1fKB/s read=%.1fKB/s\n", ...);
        return 0;
    }
    printf("usage: sd_test [block_size total_bytes]\n"); return 1;
}

/* 注册命令 → 启动 REPL */
esp_err_t console_start(void) {
    esp_console_repl_config_t repl_cfg = ESP_CONSOLE_REPL_CONFIG_DEFAULT();
    repl_cfg.prompt = "esp> ";
    esp_console_dev_uart_config_t hw = ESP_CONSOLE_DEV_UART_CONFIG_DEFAULT();
    esp_console_new_repl_uart(&hw, &repl_cfg, &repl);
    register_sd_test();  // 注册 sd_test 命令
    esp_console_start_repl(repl);
    return ESP_OK;
}'''

SRC['sd_speed_test'] = r'''/* ===== main/sd_speed_test.c (185 行) =====
 * SD 卡 FATFS 读写速度测试 — 通过 fwrite/fread 测量实际吞吐
 * 使用确定性伪随机模式填数据+校验, 测试文件自动删除 */
#include "sd_speed_test.h" "sd_card.h"
#include "esp_timer.h" "esp_heap_caps.h" "sdmmc_cmd.h"
#define TEST_PATH "/sdcard/sd_speed_test.tmp"
#define TEST_BUF_MAX (64*1024)
#define TEST_SEED 0xA55AA55Au

/* 确定性伪随机填充 (写数据+校验用同一 seed) */
static void fill_pattern(uint8_t *buf, size_t size, uint32_t s) {
    for (size_t i = 0; i < size; i++) { s = s*1103515245u+12345u; buf[i]=(uint8_t)(s>>16); }
}

/* 单次测试: 写入 N 个 block → fclose 触发 flush → 读取+校验 */
esp_err_t sd_speed_test_once(size_t block_size, size_t total_bytes, sd_speed_result_t *r) {
    size_t n = total_bytes / block_size;
    fill_pattern(s_buf, block_size, TEST_SEED); remove(TEST_PATH);
    // 写入计时
    int64_t t0 = esp_timer_get_time();
    FILE *f = fopen(TEST_PATH, "wb");
    for (size_t i = 0; i < n; i++) fwrite(s_buf, 1, block_size, f);
    fclose(f);  // fclose 触发 FATFS flush + 真正落盘
    r->write_ms = (esp_timer_get_time() - t0) / 1000;
    r->write_kbs = (double)total_bytes/1024.0 / ((double)r->write_ms/1000.0);
    // 读取+校验计时
    t0 = esp_timer_get_time();
    f = fopen(TEST_PATH, "rb");
    for (size_t i = 0; i < n; i++) { fread(s_buf, 1, block_size, f);
        if (i == 0) r->verify_ok = verify_pattern(s_buf, block_size, TEST_SEED); }
    fclose(f);
    r->read_ms = (esp_timer_get_time() - t0) / 1000;
    r->read_kbs = (double)total_bytes/1024.0 / ((double)r->read_ms/1000.0);
    remove(TEST_PATH);
    return ESP_OK;
}'''

SRC['index_html'] = r'''<!-- ===== main/index.html (309 行) =====
 Web 前端单页应用 — 纯 HTML+CSS+Vanilla JS, 零外部依赖
 200ms 增量轮询实时流 + 2s 状态刷新 + HEX 发送 + 文件管理
 通过 CMake EMBED_FILES 内嵌到固件 -->
<!DOCTYPE html><html lang=zh-CN>
<head><meta charset=UTF-8><meta name=viewport content="width=device-width,initial-scale=1">
<title>ESP Recorder</title>
<style>
body{font-family:-apple-system,sans-serif;padding:16px;max-width:640px;margin:0 auto}
.card{background:#f5f5f5;border-radius:8px;padding:14px;margin-bottom:14px}
.on{color:#0a8a00}.off{color:#888}
button{padding:9px 14px;margin:6px 6px 0 0;font-size:14px;cursor:pointer}
</style></head><body>
<h1>ESP WiFi 记录仪</h1>

<!-- 状态面板 -->
<div class=card><h2>设备状态</h2>
<div class=row><span>AP 模式</span><span id=ap>-</span></div>
<div class=row><span>STA 状态</span><span id=sta>-</span></div>
<div class=row><span>SD 卡</span><span id=sd>-</span></div>
<div class=row><span>记录状态</span><span id=rec>-</span></div>
<div class=row><span>CH1 RX/TX</span><span><b id=ch1rx>0</b>/<b id=ch1tx>0</b></span></div>
<div class=row><span>CH2 RX/TX</span><span><b id=ch2rx>0</b>/<b id=ch2tx>0</b></span></div>
<div class=row><span>路由 CH1→CH2</span><span id=route12>0</span></div>
</div>

<!-- WiFi 配网 -->
<div class=card><h2>WiFi 配网 (STA)</h2>
<label>SSID<input id=wssid placeholder=WiFi名称></label>
<select id=wauth><option value=0>开放</option><option value=3 selected>WPA2</option></select>
<label>密码<input id=wpass type=password></label>
<button onclick=saveWifi()>保存并连接</button>
</div>

<!-- 实时数据 -->
<div class=card><h2>实时数据 (TX/RX)</h2>
<select id=txchannel><option value=1>CH1(43/44)</option><option value=2>CH2(17/18)</option></select>
<textarea id=txdata placeholder="HEX: 01 03 00 0A" style="width:100%;height:60px"></textarea>
<button onclick=sendTx()>发送 TX</button>
<button onclick=clearStream()>清空显示</button>
<div id=stream style="height:260px;overflow:auto;font-family:monospace;font-size:12px">等待数据...</div>
</div>

<!-- 控制 -->
<div class=card><h2>设备控制</h2>
<button onclick=startRec()>开始记录</button>
<button onclick=stopRec()>停止记录</button>
<button onclick=reboot() style=background:#d33;color:#fff>重启设备</button>
</div>

<script>
// 实时流 200ms 增量轮询
async function pollStream(){
  const r=await fetch('/api/stream?since='+sinceSeq);const j=await r.json();
  for(const it of j.items){el.appendChild(makeLine(it.line));sinceSeq=Math.max(sinceSeq,it.seq)}
  el.scrollTop=el.scrollHeight;
  if(++pollCount>=5){trimLines();pollCount=0}
}
setInterval(pollStream,200);
// 状态 2s 刷新
async function refreshStatus(){
  const s=await (await fetch('/api/status')).json();
  document.getElementById('ap').innerHTML=s.ap?'<span class=on>已开启</span>':'<span class=off>未开启</span>';
  document.getElementById('rec').innerHTML=s.recording?'<span class=on>记录中</span>':'<span class=off>空闲</span>';
  document.getElementById('ch1rx').textContent=s.uart_ch1_rx_bytes||0;
}
setInterval(refreshStatus,2000);
// 发送 HEX
async function sendTx(){
  const r=await fetch('/api/send',{method:'POST',
    body:JSON.stringify({channel:Number(document.getElementById('txchannel').value),
    hex:document.getElementById('txdata').value.trim()})});
}
</script></body></html>'''

SRC['cmake'] = r'''# ===== 顶层 CMakeLists.txt =====
cmake_minimum_required(VERSION 3.16)
include($ENV{IDF_PATH}/tools/cmake/project.cmake)
project(esp_recorder_v1.0)

# ===== main/CMakeLists.txt =====
idf_component_register(
    SRCS "app_main.c" "console_cmd.c" "sd_speed_test.c"
         "web_server.c" "config_manager.c" "data_recorder.c"
         "data_router.c" "wifi_config.c" "usb_cdc.c"
         "usb_msc.c" "uart_driver.c" "sd_card.c"
    INCLUDE_DIRS "."
    EMBED_FILES "index.html")   # 内嵌 Web 前端到固件

# ===== main/idf_component.yml =====
dependencies:
  idf:
    version: ">=4.1.0"'''

SRC['vscode_settings'] = r'''// ===== .vscode/settings.json =====
{
  "idf.currentSetup": "D:\\23178\\esp-idf",
  "idf.customExtraVars": { "IDF_TARGET": "esp32s3" },
  "idf.flashType": "UART",
  "idf.openOcdConfigs": ["board/esp32s3-builtin.cfg"],
  "C_Cpp.default.compileCommands": "${workspaceFolder}/build/compile_commands.json",
  "clangd.arguments": [
    "--background-index",
    "--compile-commands-dir=${workspaceFolder}/build"
  ]
}'''


# ===================== Word 辅助函数 =====================

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = ''
        p = cell.paragraphs[0]; run = p.add_run(h)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]; cell.text = ''
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1: set_cell_shading(cell, 'D6E4F0')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_code(doc, code_text, label=None):
    """添加完整代码块"""
    if label:
        p = doc.add_paragraph(); run = p.add_run(label)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x2F,0x54,0x96)
        p.paragraph_format.space_after = Pt(2)
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.2)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F2F2F2'); shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = 'Consolas'; run.font.size = Pt(7.2)
    doc.add_paragraph()

def add_photo(doc, caption, w=13):
    table = doc.add_table(rows=1, cols=1); table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]; cell.width = Cm(w)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:trHeight')
    th.set(qn('w:val'), '3400'); th.set(qn('w:hRule'), 'atLeast')
    trPr.append(th)
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('📷 插入照片'); run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88,0x88,0x88); run.italic = True
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'【{caption}】'); r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x2F,0x54,0x96); r2.bold = True
    doc.add_paragraph()

def H(doc, text, level=1):
    return doc.add_heading(text, level=level)

def P(doc, text, bold=False):
    p = doc.add_paragraph(); run = p.add_run(text)
    if bold: run.bold = True
    run.font.size = Pt(10.5); return p

def B(doc, text):
    p = doc.add_paragraph(text, style='List Bullet'); return p

def page(doc):
    doc.add_page_break()


# ===================== 主构建 =====================

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)

    # ===================== 封面 =====================
    for _ in range(7): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('ESP Recorder 技术报告'); r.bold = True
    r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x2F,0x54,0x96)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('ESP32-S3 双路串口数据记录仪\n完整源代码解析与功能实现'); r2.font.size = Pt(14)
    doc.add_paragraph(); doc.add_paragraph()
    for label, val in [('项目名称','esp_recorder_v1.0'),('主控芯片','ESP32-S3 (Xtensa LX7)'),
        ('Flash','16MB'),('开发框架','ESP-IDF v5.4'),
        ('编译工具','VS Code + Espressif IDF Extension v2.1.0'),
        ('语言','C (ESP-IDF) + HTML/CSS/JavaScript (Web前端)'),
        ('总代码量','约 3,600 行 (C: ~3,300行 + Web: 309行)'),
        ('文档日期','2026-08-03'),('作者','尤译庆')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'{label}：{val}').font.size = Pt(10.5)
    page(doc)

    # ===================== 第一章 项目全景 =====================
    H(doc, '第一章 项目全景', 1)

    H(doc, '1.1 这个项目是什么', 2)
    P(doc, 'ESP Recorder 是一台基于 ESP32-S3 芯片的"串口数据黑匣子"。把它接在两台串口设备之间，它能悄无声息地记录下所有通信数据，保存到 SD 卡，同时提供一个 WiFi 热点，手机连上就能实时查看数据、下载文件、控制启停。')
    P(doc, '一句话概括：它是一个 WiFi 串口数据监听记录仪，无需电脑，自带电池就能工作。')

    H(doc, '1.2 能做什么 (16 项功能)', 2)
    add_table(doc, ['功能', '说明', '状态'], [
        ['① 双路 UART 监听', 'CH1(GPIO43/44)+CH2(GPIO17/18) 同时采集', '✅ 编译通过'],
        ['② SD 卡存储', 'SPI模式 FATFS, 自动生成 REC_CH1_时间戳.bin', '✅ 编译通过'],
        ['③ WiFi AP 热点', 'ESP_Recorder_XXXX, 192.168.4.1, 手机直连', '✅ 编译通过'],
        ['④ WiFi STA 连接', '可连外部路由器上网, WPA2/WPA3 支持', '✅ 编译通过'],
        ['⑤ Web 管理页面', '手机浏览器直接操作, 无需安装 APP', '✅ 编译通过'],
        ['⑥ 实时数据监控', 'HEX 格式, 200ms 刷新, RX/TX 双向显示', '✅ 编译通过'],
        ['⑦ 跨通道转发', 'CH1收到自动发到CH2, CH2收到自动发到CH1', '✅ 编译通过'],
        ['⑧ USB 虚拟串口', '插电脑识别为串口, 可用于调试', '✅ 编译通过'],
        ['⑨ USB 读卡器', '插电脑 SD 卡变 U 盘, 直接拷贝文件', '✅ 编译通过'],
        ['⑩ USB 热插拔保护', '插拔 USB 时自动切换 SD 卡归属权', '✅ 编译通过'],
        ['⑪ JSON 配置存储', 'SD卡 config.json, 电脑直接编辑', '✅ 编译通过'],
        ['⑫ 记录启停控制', 'Web 远程控制, 上电自动开始', '✅ 编译通过'],
        ['⑬ 双通道独立统计', '每通道独立 RX/TX/丢包/文件状态', '✅ 编译通过'],
        ['⑭ 文件管理', 'Web 列出/下载/删除 .bin 文件', '✅ 编译通过'],
        ['⑮ 调试控制台', '串口命令 sd_test 测 SD 速度', '✅ 编译通过'],
        ['⑯ 远程重启', 'Web 一键重启设备', '✅ 编译通过'],
    ], col_widths=[4, 7.5, 3])

    H(doc, '1.3 项目文件结构', 2)
    add_table(doc, ['文件', '行数', '作用', '类型'], [
        ['main/app_main.c', '183', '主入口, 13步初始化所有子系统', '入口'],
        ['main/index.html', '309', 'Web 前端 (内嵌固件)', '前端'],
        ['main/web_server.c', '649', 'HTTP 服务器 + 14个 API', '服务器'],
        ['main/data_recorder.c', '614', '双通道记录引擎 (最复杂模块)', '核心逻辑'],
        ['main/config_manager.c', '278', 'JSON 配置读写 + 字段校验', '配置'],
        ['main/wifi_config.c', '253', 'WiFi AP+STA 管理 + 事件处理', '网络'],
        ['main/uart_driver.c', '221', '双路 UART 驱动 + 字节统计', '驱动'],
        ['main/usb_msc.c', '189', 'USB MSC SCSI 读写回调', 'USB'],
        ['main/sd_speed_test.c', '185', 'SD 卡 FATFS 读写速度测试', '工具'],
        ['main/sd_card.c', '171', 'SD 卡 SPI 初始化 + FATFS 挂载', '驱动'],
        ['main/usb_cdc.c', '150', 'USB CDC 虚拟串口 + 描述符', 'USB'],
        ['main/console_cmd.c', '91', '调试控制台 REPL', '工具'],
        ['main/data_router.c', '79', '跨通道双向数据转发', '核心逻辑'],
        ['.vscode/settings.json', '19', 'VS Code ESP-IDF 工作区配置', '配置'],
    ], col_widths=[4.5, 1.5, 6.5, 2])

    H(doc, '1.4 关键技术规格', 2)
    add_table(doc, ['项目', '参数'], [
        ['芯片', 'ESP32-S3 (Xtensa LX7 双核 240MHz)'],
        ['Flash / PSRAM', '16MB / 8MB (Octal SPI)'],
        ['SD 卡接口', 'SPI2_HOST, CS=GPIO10, MOSI=38, MISO=40, SCK=39, 20MHz'],
        ['用户串口 CH1', 'UART0, TX=GPIO43, RX=GPIO44, 115200 baud, 8N1'],
        ['用户串口 CH2', 'UART2, TX=GPIO17, RX=GPIO18, 115200 baud, 8N1'],
        ['调试控制台', 'UART1, TX=GPIO1, RX=GPIO2, 115200 baud (REPL)'],
        ['USB 接口', '原生 USB OTG, CDC(虚拟串口)+MSC(读卡器) 复合设备'],
        ['USB VID/PID', '303A:4002 (Espressif)'],
        ['WiFi', '802.11 b/g/n, AP+STA 混合模式, 20MHz 带宽'],
        ['AP 默认值', 'SSID=ESP_Recorder_XXXX, 密码=12345678, IP=192.168.4.1'],
        ['Web 服务器', 'esp_http_server, 端口 80, 14 个 REST API'],
        ['文件系统', 'FATFS (SD卡), NVS (WiFi 内部存储)'],
        ['Flash 布局', 'bootloader@0x0, partition@0x8000, ota_data@0xf000, factory@0x20000 (4MB)'],
        ['Flash 参数', 'dio 模式, 80MHz, 16MB'],
        ['FreeRTOS 任务数', '5 个 (uart_rx_ch1, uart_rx_ch2, rec_io, rec_sync, tinyusb)'],
    ], col_widths=[4, 12])

    page(doc)

    # ===================== 第二章 整体架构 =====================
    H(doc, '第二章 整体架构设计', 1)

    H(doc, '2.1 系统启动流程', 2)
    P(doc, '设备上电后，按以下顺序逐步初始化每个子系统。每个步骤的失败都不影响后续步骤。')
    steps = [
        ['步骤 1', 'gpio_config()', 'GPIO16 拉低 → 使能串口接收电路'],
        ['步骤 2', 'nvs_flash_init()', '初始化 NVS 分区 (WiFi 协议栈依赖 NVS 存储)'],
        ['步骤 3', 'sd_card_init()', '初始化 SPI2 总线 → 检测 SD 卡 → 挂载 FATFS 到 /sdcard'],
        ['步骤 4', 'usb_cdc_init()', '配置 TinyUSB 复合设备描述符 → 安装驱动'],
        ['步骤 5', 'usb_msc_init()', '注册 MSC 回调 (此时 USB 复合设备完全就绪)'],
        ['步骤 6', 'config_manager_load()', '从 /sdcard/config.json 加载配置, 不存在用默认值'],
        ['步骤 7', 'recorder_init()', '创建记录引擎: 分配环形缓冲 + 启动 IO/Sync 两个 FreeRTOS 任务'],
        ['步骤 8', 'user_uart_init_channel()×2', '独立初始化 CH1/CH2 (一通道失败不阻塞另一通道)'],
        ['步骤 9', 'data_router_init()', '初始化跨通道转发'],
        ['步骤 10', 'wifi_config_init()', '初始化 WiFi 协议栈 (netif → event_loop → wifi_init)'],
        ['步骤 11', 'wifi_apply_ap()', '启动 AP 热点 (SSID/密码/IP 从 config 读取)'],
        ['步骤 12', 'web_server_start()', '启动 HTTP 服务器, 注册 14 个 API 路由'],
        ['步骤 13', 'wifi_connect_sta()', '(可选) 连接外部 WiFi'],
        ['步骤 14', 'recorder_start()', '(可选) 上电自动开始记录'],
        ['步骤 15', 'xTaskCreate(uart_rx_task)×2', '创建两路 UART 接收任务'],
        ['步骤 16', 'console_start()', '启动调试控制台 REPL'],
    ]
    add_table(doc, ['步骤', '函数调用', '具体操作'], steps, col_widths=[1.5, 4.5, 10])

    H(doc, '2.2 FreeRTOS 任务架构', 2)
    P(doc, '系统运行时有 5 个 FreeRTOS 任务同时工作：')
    add_table(doc, ['任务名', '优先级', '栈', '职责', '通信方式'], [
        ['uart_rx_ch1', '5', '4KB', 'CH1 串口接收 → 写入 recorder + router', '直接调用'],
        ['uart_rx_ch2', '5', '4KB', 'CH2 串口接收 → 写入 recorder + router', '直接调用'],
        ['rec_io', '5', '4KB', '环形缓冲落盘: 阈值≥2KB或超时1秒 → fwrite', '信号量 (s_data_sem)'],
        ['rec_sync', '1', '3KB', '每秒 fsync() 强制同步, 连续3次失败→关闭文件', '定时器 (vTaskDelay)'],
        ['tinyusb', '5', '4KB', 'TinyUSB 协议栈 (CPU1亲和, CDC+MSC 复合)', '中断 → 回调'],
    ], col_widths=[2.5, 1.5, 1.5, 6.5, 4])

    H(doc, '2.3 完整数据流', 2)
    P(doc, '下面用一张文字流程图展示数据如何在系统中流转：')
    flow = (
        '━━━━━━━━━━━━━━━ 数据流全景图 ━━━━━━━━━━━━━━━\n\n'
        '外部设备A TX ──→ CH1 UART (RX44) ──→ uart_rx_ch1 任务\n'
        '    │\n'
        '    ├──→ recorder_write_rx_channel(CH1, data, len)\n'
        '    │      ├── live_push() → 实时流环形缓冲 (256条)\n'
        '    │      │      └── Web /api/stream?since=N 增量拉取 → 手机显示\n'
        '    │      └── ringbuf_write() → 落盘环形缓冲 (16KB/通道)\n'
        '    │             └── rec_io 任务 (阈值/超时触发)\n'
        '    │                    └── fwrite() → /sdcard/REC_CH1_时间戳.bin\n'
        '    │\n'
        '    └──→ data_router_forward(CH1→CH2)\n'
        '           └── user_uart_write_channel(CH2) → TX17 发出\n'
        '                  └── 外部设备B 收到\n\n'
        '外部设备B TX ──→ CH2 UART (RX18) ──→ uart_rx_ch2 任务\n'
        '    ├──→ recorder_write_rx_channel(CH2) → SD卡 REC_CH2_*.bin\n'
        '    └──→ data_router_forward(CH2→CH1) → TX43 发出 → 外部设备A\n\n'
        'USB 插入:\n'
        '  tud_mount_cb() → recorder_stop() → sd_card_unmount_fs()\n'
        '                 → SD 卡交给 USB MSC → 电脑弹出 U 盘\n\n'
        'USB 拔出:\n'
        '  tud_umount_cb() → sd_card_remount_fs() → recorder_start()\n'
        '                 → 恢复 SD 卡记录\n\n'
        'Web 操作:\n'
        '  POST /api/send {\"channel\":1,\"hex\":\"01 03\"} → TX 发出\n'
        '  GET  /api/stream?since=N → 增量实时流 JSON')
    add_code(doc, flow, '▼ 数据流全景')

    page(doc)

    # ===================== 第三章 逐模块源码解析 =====================
    H(doc, '第三章 逐模块源码解析', 1)
    P(doc, '本章对项目 14 个源文件逐一进行完整源码展示 + 逐行拆解分析。每个模块都说明了：① 代码做了什么 ② 为什么这样设计 ③ 功能是怎么一步步实现的。')

    # ---- 3.1 主入口 ----
    H(doc, '3.1 主入口 — app_main.c (183行)', 2)
    P(doc, '这是整个系统的启动核心。上电后进入 app_main() 函数，按照严格的顺序初始化全部 13 个子系统。')
    P(doc, '设计要点：')
    B(doc, 'GPIO16 在最先操作 — 因为它是串口接收的使能引脚，必须在串口使用前拉低')
    B(doc, 'SD 卡先于配置加载 — 因为 config.json 在 SD 卡上，必须先挂载 SD 才能读配置')
    B(doc, 'USB 先于 WiFi — 让 USB 复合设备尽快暴露，即使后面 WiFi 初始化失败仍可 USB 调试')
    B(doc, 'UART 独立初始化 — 两路各自 try，一路失败不会让另一路罢工')
    B(doc, 'Web 服务器在 WiFi 之后 — Web 依赖 WiFi AP 的 IP 地址存在')
    B(doc, '可选步骤都有条件判断 — wifi_connect_sta() 和 recorder_start() 根据配置决定')
    add_code(doc, SRC['app_main'], '▼ 完整源码: app_main.c')

    P(doc, '逐行拆解：')
    P(doc, '第 1-8 行: 头文件包含。项目所有模块的头文件在这里集中引入, 体现了模块间的依赖关系。')
    P(doc, '第 24-28 行: usb_cdc_rx_handler() 是 USB 虚拟串口收到数据时的回调。目前为空实现, 保留了扩展接口——未来可以在这里加 USB 命令处理。')
    P(doc, '第 32-37 行: tud_mount_cb() 是 USB 插入回调。核心逻辑是先停止记录（防止 SD 卡被 FATFS 和 USB MSC 同时访问），再卸载文件系统。')
    P(doc, '第 40-51 行: tud_umount_cb() 是 USB 拔出回调。重新挂载 FATFS 后, 如果配置了 auto_start, 自动恢复记录。')
    P(doc, '第 53-68 行: uart_rx_task() 是每个 UART 通道的接收任务函数。用一个任务函数服务两个通道: 通过 FreeRTOS 任务参数传递通道号。100ms 超时读取避免死等。')
    P(doc, '第 71-183 行: app_main() 是真正的入口。13 个步骤按顺序执行, 每一步的错误处理都是 ESP_ERROR_CHECK 或单独检查。')

    # ---- 3.2 UART 驱动 ----
    H(doc, '3.2 UART 驱动 — uart_driver.h + uart_driver.c (268行)', 2)
    P(doc, '这是串口通信的底层驱动。每个通道有独立的 UART 编号、TX 引脚、RX 引脚、缓冲区。')
    P(doc, '设计要点：')
    B(doc, '引脚定义在 .h 中用宏声明 — 方便改硬件时只需改一处')
    B(doc, '每通道独立初始化 — app_main 中 for 循环调用, 一个失败不影响另一个')
    B(doc, 'portMUX_TYPE 自旋锁 — 保证多任务访问 RX/TX 计数器时的线程安全')
    B(doc, 'TX 回调机制 — 每次成功发送后回调 recorder_write_tx_channel(), 把发送数据也推送到 Web 实时流')
    B(doc, '首次 vs 后续初始化区分 — uart_is_driver_installed() 判断是首次安装驱动还是仅修改波特率')
    add_code(doc, SRC['uart_h'], '▼ 完整源码: uart_driver.h (引脚定义与接口)')
    add_code(doc, SRC['uart_c'], '▼ 完整源码: uart_driver.c (初始化 + 读写 + 统计)')

    # ---- 3.3 SD 卡驱动 ----
    H(doc, '3.3 SD 卡驱动 — sd_card.c (171行)', 2)
    P(doc, '负责 SD 卡的物理访问和文件系统挂载。采用 SPI 模式而非 SDMMC 模式，释放 SDMMC 引脚给其他用途。')
    P(doc, '设计要点：')
    B(doc, 'SPI 总线只初始化一次 (s_spi_initialized 标志) — 避免重复初始化破坏已有状态')
    B(doc, '频率限制 20MHz — SPI 模式下过高频率会导致 SD 卡通信不稳定')
    B(doc, '不自动格式化 (format_if_mount_failed=false) — 保护用户 SD 卡上的已有数据')
    B(doc, 'sd_card_unmount_fs() 的特殊设计 — 只卸载 VFS/FATFS, 保留 SPI 总线和 sdmmc_card_t, 让 USB MSC 可以继续访问块设备')
    B(doc, 'sd_card_remount_fs() 是完整重挂载 — 先完全释放再重新初始化, 因为 USB 已释放硬件')
    add_code(doc, SRC['sd_card'], '▼ 完整源码: sd_card.c')

    # ---- 3.4 USB 复合设备 ----
    H(doc, '3.4 USB 复合设备 — usb_cdc.c + usb_msc.c (339行)', 2)
    P(doc, '让 ESP32-S3 的原生 USB 接口同时提供两个功能：虚拟串口(CDC) + SD 读卡器(MSC)。一根 USB 线, 既能调试又能传文件。')
    P(doc, '设计要点：')
    B(doc, 'USB 描述符是复合设备的关键 — desc_configuration[] 中 CDC 和 MSC 接口并列, 电脑会识别为一个设备但有两个功能')
    B(doc, 'VID=303A 是 Espressif 官方 VID — 设备管理器中显示为 Espressif 设备')
    B(doc, 'MSC 读写带 3 次重试 — SD 卡在 SPI 模式下偶发通信错误, 重试机制大幅提高可靠性')
    B(doc, 'VBUS 检测用 GPIO48 — 这是 ESP32-S3 USB OTG 标准的 VBUS 引脚')
    add_code(doc, SRC['usb_cdc'], '▼ 完整源码: usb_cdc.c (CDC虚拟串口 + USB描述符)')
    add_code(doc, SRC['usb_msc'], '▼ 完整源码: usb_msc.c (MSC 读卡器 + SCSI 命令)')

    # ---- 3.5 WiFi 管理 ----
    H(doc, '3.5 WiFi 管理 — wifi_config.c (253行)', 2)
    P(doc, '管理 AP 热点(手机连设备) 和 STA(设备连路由器) 两种 WiFi 模式。采用 AP+STA 混合模式, 两者可以同时工作。')
    P(doc, '设计要点：')
    B(doc, 'AP 永远在线 — 即使 STA 断线, AP 不受影响, 保证用户永远可以通过热点访问 Web 管理页面')
    B(doc, 'STA 断线自动重连 — 最多重试 3 次, 间隔 5 秒, 全部失败后停止重试但不影响 AP')
    B(doc, '默认 SSID 包含 MAC 地址 — ESP_Recorder_XXYY 保证多台设备不重名')
    B(doc, 'AP 密码空串 = 开放热点 — 方便演示和调试, 生产环境可设密码')
    add_code(doc, SRC['wifi_config'], '▼ 完整源码: wifi_config.c')

    page(doc)

    # ---- 3.6 记录引擎 ----
    H(doc, '3.6 数据记录引擎 — data_recorder.c (614行) ★最核心', 2)
    P(doc, '这是全项目最复杂的模块, 也是"记录仪"这个产品的核心价值所在。它实现了双通道独立记录、环形缓冲、双条件落盘、数据安全保证。')
    add_code(doc, SRC['data_recorder'], '▼ 完整源码: data_recorder.c (核心逻辑)')

    P(doc, '逐层拆解——这个记录引擎是怎么一步步设计出来的：', bold=True)
    P(doc, '')
    P(doc, '第一层: 环形缓冲区 (Ring Buffer)', bold=True)
    P(doc, '为什么不直接 fwrite? 因为 SD 卡每次写入都有固定开销(文件系统元数据更新、Flash 擦写)。如果每收到一个字节就写一次 SD 卡, 不仅慢还会快速磨损 SD 卡。环形缓冲的作用是把零散数据攒起来, 攒够一批再写。16KB 是实验值——太大浪费内存, 太小写入次数多。')
    P(doc, '')
    P(doc, '第二层: 双条件触发落盘', bold=True)
    P(doc, '只有环形缓冲还不够。如果数据来得慢, 比如一分钟才来几个字节, 缓冲永远攒不满, 数据就永远留在内存里, 断电就丢了。所以加了"超时 1 秒"的兜底条件。两个条件满足任意一个就触发 fwrite: ① 缓冲 ≥ 2KB ② 距离上次写入超过 1 秒。')
    P(doc, '')
    P(doc, '第三层: fsync 保证数据安全', bold=True)
    P(doc, 'fwrite 只是把数据写到操作系统的缓冲区(在 FATFS 层), 还没真正落到 SD 卡上。如果此时断电, 数据会丢失。fsync() 强制把 FATFS 缓冲区的内容写入 SD 卡的物理扇区。每秒调用一次, 确保最多丢失 1 秒的数据。连续 3 次 fsync 失败说明 SD 卡可能坏了或被拔了, 主动关闭文件防止数据错乱。')
    P(doc, '')
    P(doc, '第四层: 实时流机制', bold=True)
    P(doc, '除了落盘, 数据还要送到 Web 前端实时显示。实时流和落盘用不同的缓冲区(虽然内容相同): 实时流用 256 条环形缓冲, 每条包含时间戳+通道+方向+HEX数据。Web 端 200ms 轮询一次, 用 since_seq 机制实现增量拉取——只取上次序号之后的新数据, 避免重复传输。')
    P(doc, '')
    P(doc, '第五层: 双通道独立', bold=True)
    P(doc, '每个通道有自己的环形缓冲、文件名、统计计数器。CH1 写满了不影响 CH2 写, CH1 文件关闭了不影响 CH2 继续记录。通道之间完全解耦。')
    P(doc, '')
    P(doc, '第六层: 文件命名防冲突', bold=True)
    P(doc, '文件名用时间戳生成: REC_CH1_20260727_193428.bin。如果在同一秒内多次开始记录(极少见), 自动加 _1, _2… 后缀, 循环找到第一个不存在的文件名。')

    # ---- 3.7 数据路由 ----
    H(doc, '3.7 数据路由 — data_router.c (79行)', 2)
    P(doc, '这个模块实现了一个简单的功能: 把 CH1 收到的数据转发到 CH2 TX, 把 CH2 收到的数据转发到 CH1 TX。相当于在硬件层面把两个串口设备"打通"了。')
    P(doc, '使用场景: 设备A 接 CH1, 设备B 接 CH2。设备A 发数据给设备B 时, 路由器自动转发, 同时 recorder 全程记录。这样两个设备以为它们在直接通信, 实际上中间有一个透明的记录仪。')
    add_code(doc, SRC['data_router'], '▼ 完整源码: data_router.c')

    # ---- 3.8 Web 服务器 ----
    H(doc, '3.8 Web 服务器 — web_server.c (649行)', 2)
    P(doc, '基于 ESP-IDF 内置的 esp_http_server 组件实现。它是一个轻量级的 HTTP 服务器, 监听 80 端口, 提供了 14 个 REST API 端点。')
    add_code(doc, SRC['web_server'], '▼ 完整源码: web_server.c (路由表 + 核心 Handler)')

    P(doc, 'API 设计解析：')
    P(doc, 'GET /api/status — 返回完整设备状态 JSON, 含 25+ 个字段。前端每 2 秒调用一次, 更新所有面板数据。使用 cJSON 库而非 snprintf 拼接——因为 SSID 可能包含引号、反斜杠等特殊字符, 直接拼接会损坏 JSON。')
    P(doc, 'POST /api/wifi — 接收前端提交的 SSID/密码, 写入 config.json 持久化, 然后调用 wifi_connect_sta() 触发连接。即使 WiFi 连接失败, 配置也已被保存, 下次上电自动重试。')
    P(doc, 'POST /api/send — 把用户输入的 HEX 字符串 (如 "01 03 00 0A") 解码为二进制字节, 通过 UART TX 发出去。同时回调 recorder_write_tx_channel() 把发送的数据也显示在实时流中。')
    P(doc, 'GET /api/stream?since=N — 实时流增量拉取。since=N 告诉服务器"我已经拿到了序号 N 的数据, 只要 N 之后的"。前端记录最新收到的序号, 每次请求只传 since_seq, 服务端只返回新数据。')
    P(doc, 'GET /api/file?name=REC_CH1_xxx.bin — 文件下载。设置 Content-Disposition: attachment 头, 浏览器自动弹出另存为对话框。使用分块传输 (chunked) 而非一次加载整个文件, 避免大文件耗尽内存。')

    # ---- 3.9 配置管理器 ----
    H(doc, '3.9 配置管理器 — config_manager.c (278行)', 2)
    P(doc, '负责读取和写入 /sdcard/config.json。配置放在 SD 卡而不是 NVS 中, 是为了方便用户直接用电脑编辑 JSON 文件, 不需要任何专用工具。')
    P(doc, '安全设计：')
    B(doc, '加载时先填默认值, 再逐字段覆盖 — 缺失字段保留默认值, 旧版配置文件也能正常工作')
    B(doc, 'authmode 合法性校验 — config 中如果写了非法值(如 99), clamp 到合法值(0=OPEN)')
    B(doc, 'stream_gap_ms 限制范围 [1, 1000] — 防止用户输入 0 或超大值导致异常')
    B(doc, '保存失败自动回滚 — 如果写入 SD 卡失败, 内存中的值回退到旧值, 不会留下一半修改的状态')
    add_code(doc, SRC['config_manager'], '▼ 完整源码: config_manager.c')

    # ---- 3.10 调试控制台 ----
    H(doc, '3.10 调试控制台 + SD 速度测试 — console_cmd.c + sd_speed_test.c (276行)', 2)
    P(doc, '调试控制台使用 ESP-IDF 的 Console REPL 组件, 提供交互式命令行。目前只注册了一个 sd_test 命令, 但架构支持未来添加更多调试命令。')
    P(doc, 'SD 速度测试不是玩具——它通过 fwrite/fread 测量 FATFS 层的实际读写吞吐量, 完全模拟真实记录场景。先用确定性伪随机模式填充数据, 写入后 fclose 触发 flush, 再读回校验。')
    add_code(doc, SRC['console'], '▼ 完整源码: console_cmd.c')
    add_code(doc, SRC['sd_speed_test'], '▼ 完整源码: sd_speed_test.c')

    # ---- 3.11 Web 前端 ----
    H(doc, '3.11 Web 前端 — index.html (309行)', 2)
    P(doc, '这是一个纯 HTML+CSS+Vanilla JavaScript 的单页应用, 零外部依赖。固件编译时通过 CMake 的 EMBED_FILES 指令把 index.html 嵌入到固件中, 所以即使 SD 卡上没有 Web 文件, 网页也能正常打开。')
    P(doc, '前端核心 JS 函数：')
    B(doc, 'pollStream() — 每 200ms 调用 /api/stream?since=N, 增量获取实时流数据。超过 stream_gap_ms 毫秒没收到新数据时, 插入"等待数据..."标记')
    B(doc, 'refreshStatus() — 每 2 秒调用 /api/status, 更新所有面板数据')
    B(doc, 'sendTx() — 读取文本框中的 HEX 字符串, POST 到 /api/send')
    B(doc, 'trimLines() — 限制最多显示 500 行, 超出后截断前部, 防止 DOM 节点过多导致浏览器卡顿')
    add_code(doc, SRC['index_html'], '▼ 完整源码: index.html')

    # ---- 3.12 构建配置 ----
    H(doc, '3.12 构建配置 — CMakeLists.txt + .vscode', 2)
    P(doc, 'CMakeLists.txt 定义了哪些源文件参与编译。EMBED_FILES "index.html" 是关键——它把 HTML 文件转为 C 数组链接进固件。')
    add_code(doc, SRC['cmake'], '▼ 完整源码: CMakeLists.txt + idf_component.yml')
    add_code(doc, SRC['vscode_settings'], '▼ 完整源码: .vscode/settings.json')

    page(doc)

    # ===================== 第四章 关键技术决策 =====================
    H(doc, '第四章 关键技术决策与设计理由', 1)

    add_table(doc, ['序号', '决策', '选择的方案', '为什么不选其他方案'], [
        ['1', 'SD卡接口', 'SPI 模式', '释放 SDMMC 引脚给其他外设; SPI 速率(~2MB/s)足够串口记录使用'],
        ['2', '配置存储位置', 'SD 卡 JSON', '用户可用电脑直接编辑, 无需专用工具; NVS 需要写代码读写'],
        ['3', 'WiFi 模式', 'AP+STA 混合', 'AP 永远可用保证本地管理; STA 可选连外网; 两者互不阻塞'],
        ['4', '双通道初始化策略', '独立初始化', '某通道硬件故障不影响另一通道; 系统级健壮性'],
        ['5', 'USB 方案', 'TinyUSB 复合设备', '一根线同时提供调试+取文件; 比分立 USB-UART 芯片更简洁'],
        ['6', 'JSON 拼装', 'cJSON 库', 'SSID 含特殊字符时 snprintf 会损坏 JSON; cJSON 自动转义'],
        ['7', 'Web 部署', 'EMBED_FILES 内嵌', '固件自包含, 不依赖 SD 卡; 即使卡坏了网页也能打开'],
        ['8', '落盘触发条件', '双条件 (阈值/超时)', '数据快时批量写减少擦写; 数据慢时超时兜底保证不丢'],
        ['9', '数据安全', '每秒 fsync()', '最多丢1秒数据; 连续3次失败主动关闭文件防错乱'],
        ['10', '文件格式', 'Raw 二进制 .bin', '不添加文件头/时间戳, 保持原始字节流; 用户可以自行解析'],
        ['11', '实时流方式', '增量轮询 since_seq', '只传新数据, 不重复; 比 WebSocket 更简单可靠'],
        ['12', 'USB 热插拔', 'FATFS/MSC 切换', 'USB 插入时保证 SD 卡不被两边同时写; 拔出自动恢复'],
    ], col_widths=[0.8, 3, 4.5, 7.5])

    page(doc)

    # ===================== 第五章 硬件接线 =====================
    H(doc, '第五章 硬件接线图', 1)

    H(doc, '5.1 SD 卡 (SPI)', 2)
    add_table(doc, ['SD 卡引脚', 'ESP32-S3'], [['CS','GPIO10'],['MOSI','GPIO38'],['MISO','GPIO40'],['SCK','GPIO39'],['VCC','3.3V'],['GND','GND']], col_widths=[6,6])

    H(doc, '5.2 用户串口', 2)
    add_table(doc, ['通道','UART','TX','RX','波特率'], [['CH1','UART0','GPIO43','GPIO44','115200'],['CH2','UART2','GPIO17','GPIO18','115200']], col_widths=[2.5,2.5,3,3,4])

    H(doc, '5.3 其他引脚', 2)
    add_table(doc, ['功能','GPIO'], [['UART 使能(低有效)','16'],['USB VBUS 检测','48'],['调试控制台 TX','1'],['调试控制台 RX','2']], col_widths=[6,6])

    add_photo(doc, '图5-1: ESP32-S3 完整接线照片')
    add_photo(doc, '图5-2: SD 卡模块接线特写')

    page(doc)

    # ===================== 第六章 构建编译 =====================
    H(doc, '第六章 构建与烧录', 1)

    H(doc, '6.1 环境准备', 2)
    add_code(doc, '''# 1. 安装 ESP-IDF v5.4 (推荐 VS Code 扩展自动安装)
# 2. 安装 VS Code 扩展: espressif.esp-idf-extension v2.1.0
# 3. 打开项目
cd D:\\zhuomian\\weite\\ESP32\\esp_recorder_v1.1
code .
# 4. 编译
idf.py build''')

    H(doc, '6.2 Flash 布局', 2)
    add_table(doc, ['偏移', '分区', '大小', '内容'], [
        ['0x0', 'bootloader', '32KB', '二级引导程序'],
        ['0x8000', 'partition_table', '4KB', '分区表'],
        ['0xf000', 'ota_data', '8KB', 'OTA状态数据'],
        ['0x20000', 'factory', '4MB', 'esp_recorder_v1.0.bin (主程序)'],
    ], col_widths=[2.5, 3.5, 2.5, 7.5])

    H(doc, '6.3 烧录', 2)
    P(doc, '⚠️ 烧录前务必通过设备管理器确认 COM 口是 ESP32-S3 真实串口（VID=303A），不是蓝牙虚拟串口。')
    add_code(doc, '''idf.py -p COMx flash              # 烧录
idf.py -p COMx flash monitor      # 烧录 + 看日志''')

    add_photo(doc, '图6-1: VS Code 终端 idf.py build 构建成功')
    add_photo(doc, '图6-2: 烧录过程终端截图')

    page(doc)

    # ===================== 第七章 上电使用 =====================
    H(doc, '第七章 上电使用步骤', 1)

    P(doc, '第一步: 插上 SD 卡, 上电。设备自动启动 AP 热点。')
    P(doc, '第二步: 手机搜索 WiFi "ESP_Recorder_XXXX", 密码 12345678, 连接。')
    P(doc, '第三步: 浏览器打开 http://192.168.4.1，即可看到 Web 管理界面。')
    P(doc, '第四步: 将外部设备串口线接到 CH1/CH2，数据自动开始记录（如果 auto_start=1）。')
    P(doc, '第五步: 在 Web 页面上可以实时看数据、下载文件、或停止记录后把 SD 卡拔出来用电脑读。')
    P(doc, '或者: 用 USB 线连接电脑 USB OTG 口，SD 卡直接在电脑上显示为 U 盘，直接拷贝文件。')

    add_photo(doc, '图7-1: 手机连接 ESP_Recorder_XXXX WiFi 热点')
    add_photo(doc, '图7-2: 浏览器打开 192.168.4.1 Web 首页全貌')
    add_photo(doc, '图7-3: 设备状态面板 (AP已开启/SD已挂载/记录中)')
    add_photo(doc, '图7-4: AP 热点配置界面')
    add_photo(doc, '图7-5: WiFi 配网 (STA) 界面')
    add_photo(doc, '图7-6: 实时数据监控 — HEX 流动态滚动')
    add_photo(doc, '图7-7: 文件列表 — REC_CH1_*.bin / REC_CH2_*.bin')
    add_photo(doc, '图7-8: 发送 HEX 到指定通道')
    add_photo(doc, '图7-9: 设备控制按钮 (开始/停止/重启)')
    add_photo(doc, '图7-10: 设备管理器 — COM 口列表 + 磁盘驱动器 (USB MSC U盘)')
    add_photo(doc, '图7-11: 串口调试控制台 esp> sd_test 输出')
    add_photo(doc, '图7-12: SD 卡中 REC_CH1_*.bin 文件在电脑上的截图')

    page(doc)

    # ===================== 第八章 常见问题 =====================
    H(doc, '第八章 常见问题排查', 1)
    add_table(doc, ['现象', '原因', '解决'], [
        ['Write timeout 烧录失败', 'COM 口是蓝牙虚拟串口', '设备管理器确认 VID=303A 的真实 COM'],
        ['SD "not mounted"', '未插卡或 SPI 接线错', '查卡座焊接, MOSI/MISO 是否交叉'],
        ['WiFi 热点搜不到', 'VBUS 检测异常', '检查 USB 连接, 确认 GPIO48 配置'],
        ['记录文件 0 字节', '串口无数据或波特率错', '检查 TX→RX 交叉接线, 确认 115200'],
        ['Web 页面白屏', 'WiFi 断开或 IP 变化', '重连热点, 刷新 http://192.168.4.1'],
    ], col_widths=[4.5, 5, 6.5])

    page(doc)

    # ===================== 附录 =====================
    H(doc, '附录 A: 完整配置文件 (config.json)', 1)
    add_code(doc, '''{
  "uart":     { "baudrate": 115200, "databits": 8, "stopbits": 1, "parity": 0 },
  "wifi":     { "enable_ap": 1, "ap_ssid": "", "ap_pass": "12345678",
                "ap_ip": "192.168.4.1", "enable_sta": 0, "sta_ssid": "",
                "sta_pass": "", "sta_authmode": 0 },
  "recorder": { "auto_start": 1, "stream_gap_ms": 100 }
}''')

    H(doc, '附录 B: 源文件完整清单', 1)
    add_table(doc, ['序号', '文件', '行数', '模块类型', '核心功能'], [
        ['1', 'main/app_main.c', '183', '入口', '系统启动13步初始化 + USB热插拔回调'],
        ['2', 'main/web_server.c', '649', '服务器', 'HTTP 80端口 + 14个 REST API'],
        ['3', 'main/data_recorder.c', '614', '核心', '双通道记录 + 环形缓冲 + 双条件落盘 + fsync'],
        ['4', 'main/config_manager.c', '278', '配置', 'JSON 加载/保存 + 字段校验 + 回滚'],
        ['5', 'main/wifi_config.c', '253', '网络', 'WiFi AP+STA + 事件处理 + 自动重连'],
        ['6', 'main/uart_driver.c', '221', '驱动', '双路 UART + RX/TX 统计 + 回调'],
        ['7', 'main/usb_msc.c', '189', 'USB', 'MSC SCSI READ10/WRITE10 + 3次重试'],
        ['8', 'main/sd_card.c', '171', '驱动', 'SD SPI + FATFS + 卸载/重挂载'],
        ['9', 'main/usb_cdc.c', '150', 'USB', 'CDC 虚拟串口 + 复合设备描述符'],
        ['10', 'main/sd_speed_test.c', '185', '工具', 'SD 卡 FATFS 读写测速'],
        ['11', 'main/console_cmd.c', '91', '工具', '调试 REPL + sd_test 命令'],
        ['12', 'main/data_router.c', '79', '核心', 'CH1↔CH2 双向数据转发'],
        ['13', 'main/index.html', '309', '前端', 'Web SPA (零依赖)'],
    ], col_widths=[0.8, 4, 1.2, 1.5, 7])

    H(doc, '附录 C: FreeRTOS 任务完整参数', 1)
    add_table(doc, ['任务名', '优先级', '栈', '创建位置', '循环逻辑', '退出条件'], [
        ['uart_rx_ch1', '5', '4096', 'app_main()', 'while(1): uart_read(100ms超时)→recorder+router', '永不退出'],
        ['uart_rx_ch2', '5', '4096', 'app_main()', 'while(1): uart_read(100ms超时)→recorder+router', '永不退出'],
        ['rec_io', '5', '4096', 'recorder_init()', '等信号量(有新数据/超时1s)→检查阈值→flush', '永不退出'],
        ['rec_sync', '1', '3072', 'recorder_init()', 'vTaskDelay(1s)→fsync(每通道)→连续3次失败关文件', '永不退出'],
        ['tinyusb', '5', '4096', 'TinyUSB 组件', 'tusb_device_task()→处理USB事件/CDC数据/MSC命令', '永不退出'],
    ], col_widths=[2.5, 1.5, 1.5, 2.5, 6, 2])

    H(doc, '附录 D: 环境恢复命令', 1)
    add_code(doc, '''idf.py --version                    # 确认 ESP-IDF v5.4
python --version                    # 确认 Python 3.8+
code --list-extensions | grep espressif    # 确认扩展
idf.py build                        # 编译验证''')

    # ===================== 保存 =====================
    out = r'D:\zhuomian\weite\ESP32\ESP_Recorder_技术报告_完整源码解析.docx'
    doc.save(out)
    print(f'✅ 已生成: {out}')
    return out

if __name__ == '__main__':
    build()
