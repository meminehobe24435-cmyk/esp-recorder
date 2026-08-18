# -*- coding: utf-8 -*-
"""将 ESP_Recorder 项目报告转换为带真实代码片段的格式化 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ====================== 真实项目代码片段 ======================

CODE_app_main = r'''#include "sd_card.h" "uart_driver.h" "usb_cdc.h" "usb_msc.h"
#include "wifi_config.h" "data_recorder.h" "data_router.h"
#include "config_manager.h" "web_server.h" "console_cmd.h"

void app_main(void) {
    // 1. GPIO16 拉低使能串口接收
    gpio_config_t uart_select = {
        .pin_bit_mask = 1 << GPIO_NUM_16,
        .mode = GPIO_MODE_OUTPUT, ... };
    gpio_config(&uart_select);
    gpio_set_level(GPIO_NUM_16, 0);

    // 2. 初始化 NVS (WiFi 需要)
    nvs_flash_init();
    // 3. SD 卡挂载 → FATFS
    sd_card_init();
    // 4. USB 复合设备 (CDC + MSC)
    usb_cdc_init(usb_cdc_rx_handler);
    usb_msc_init();
    // 5. 加载 /sdcard/config.json (不存在则创建默认)
    config_manager_t *cfg_mgr = config_manager_instance();
    config_manager_load(cfg_mgr, "/sdcard/config.json");
    const device_config_t *cfg = config_manager_get(cfg_mgr);
    // 6. 记录引擎初始化 (环形缓冲 + IO/Sync 任务)
    recorder_init();
    // 7. 双路 UART 独立初始化 (一通道失败不影响另一通道)
    for (uint8_t ch = USER_UART_CHANNEL_1; ch <= USER_UART_CHANNEL_2; ch++)
        user_uart_init_channel(ch, cfg->uart.baudrate);
    user_uart_set_tx_callback(recorder_write_tx_channel);
    data_router_init();
    // 8. WiFi AP 热点
    wifi_config_init();
    wifi_apply_ap(cfg->wifi.ap_ssid, cfg->wifi.ap_pass, cfg->wifi.ap_ip);
    // 9. HTTP Web 服务器 (端口 80)
    web_server_start();
    // 10. 可选: 连接外部 WiFi (STA 模式)
    if (cfg->wifi.enable_sta && cfg->wifi.sta_ssid[0])
        wifi_connect_sta(cfg->wifi.sta_ssid, cfg->wifi.sta_pass,
                         cfg->wifi.sta_authmode);
    // 11. 可选: 上电自动开始记录
    if (cfg->recorder.auto_start) recorder_start();
    // 12. 创建两路 UART RX 任务
    for (uint8_t ch = USER_UART_CHANNEL_1; ch <= USER_UART_CHANNEL_2; ch++)
        xTaskCreate(uart_rx_task, "uart_rx", 4096, (void*)ch, 5, NULL);
    // 13. 调试控制台
    console_start();
}'''

CODE_full_init_flow = r'''app_main() 完整初始化流程 (13 步)：
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
① gpio_config(GPIO_NUM_16)      → 拉低 GPIO16 使能串口接收
② nvs_flash_init()              → NVS 初始化 (WiFi 必需)
③ sd_card_init()                → SD 卡 SPI 挂载 FATFS
④ usb_cdc_init() + usb_msc_init() → USB 复合设备初始化
⑤ config_manager_load()         → 加载 /sdcard/config.json
⑥ recorder_init()               → 记录引擎 (环形缓冲+IO/Sync任务)
⑦ user_uart_init_channel() ×2   → CH1/CH2 独立初始化
⑧ data_router_init()            → 跨通道路由初始化
⑨ wifi_config_init()            → WiFi 协议栈初始化
⑩ wifi_apply_ap()               → 启动 AP 热点
⑪ web_server_start()            → HTTP 服务器启动 (端口80)
⑫ wifi_connect_sta()            → 可选 STA 外部 WiFi 连接
⑬ xTaskCreate(uart_rx_task) ×2  → 创建 UART RX 接收任务
    console_start()              → 调试控制台启动
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'''

CODE_uart_h = r'''// ===== uart_driver.h — 双路 UART 引脚与接口定义 =====
#define USER_UART_CHANNEL_COUNT 2
#define USER_UART_CHANNEL_1     1
#define USER_UART_CHANNEL_2     2
#define USER_UART1_NUM          UART_NUM_0
#define USER_UART1_TX_PIN       43   // CH1 发送引脚
#define USER_UART1_RX_PIN       44   // CH1 接收引脚
#define USER_UART2_NUM          UART_NUM_2
#define USER_UART2_TX_PIN       17   // CH2 发送引脚
#define USER_UART2_RX_PIN       18   // CH2 接收引脚
#define USER_UART_BUF_SIZE      2048

esp_err_t user_uart_init_channel(uint8_t channel, int baud_rate);
int user_uart_write_channel(uint8_t channel, const uint8_t *data, size_t len);
int user_uart_read_channel(uint8_t channel, uint8_t *buf, size_t len,
                           uint32_t timeout_ms);
bool user_uart_channel_is_ready(uint8_t channel);
uint64_t user_uart_get_rx_bytes(uint8_t channel);
uint64_t user_uart_get_tx_bytes(uint8_t channel);'''

CODE_uart_init = r'''// ===== uart_driver.c — 单通道 UART 初始化 =====
esp_err_t user_uart_init_channel(uint8_t channel, int baud_rate) {
    const user_uart_hw_t *hw = user_uart_get_hw(channel);
    if (hw == NULL || baud_rate <= 0) return ESP_ERR_INVALID_ARG;

    if (!uart_is_driver_installed(hw->port)) {
        uart_config_t uart_config = {
            .baud_rate = baud_rate,
            .data_bits = UART_DATA_8_BITS,
            .parity    = UART_PARITY_DISABLE,
            .stop_bits = UART_STOP_BITS_1,
            .flow_ctrl = UART_HW_FLOWCTRL_DISABLE,
            .source_clk = UART_SCLK_DEFAULT,
        };
        uart_param_config(hw->port, &uart_config);
        uart_set_pin(hw->port, hw->tx_pin, hw->rx_pin,
                     UART_PIN_NO_CHANGE, UART_PIN_NO_CHANGE);
        uart_driver_install(hw->port, USER_UART_BUF_SIZE,
                           USER_UART_BUF_SIZE, 0, NULL, 0);
    }
    s_channel_ready[index] = true;
    return ESP_OK;
}'''

CODE_uart_rx_task = r'''// ===== app_main.c — UART 接收任务 (每通道一个) =====
static void uart_rx_task(void *arg) {
    uint8_t channel = (uint8_t)(uintptr_t)arg;
    uint8_t buf[256];
    while (1) {
        int len = user_uart_read_channel(channel, buf, sizeof(buf), 100);
        if (len > 0) {
            // ① 写入记录引擎 (环形缓冲 → SD 卡落盘)
            recorder_write_rx_channel(channel, buf, len);
            // ② 数据路由 (CH1→CH2, CH2→CH1 双向转发)
            data_router_forward(channel, buf, (size_t)len);
        }
    }
}'''

CODE_sd_spi = r'''// ===== sd_card.c — SD 卡 SPI 初始化 =====
#define SD_MOUNT_POINT      "/sdcard"
#define SD_HOST             SPI2_HOST
#define SD_PIN_CS           10
#define SD_PIN_MOSI         38
#define SD_PIN_MISO         40
#define SD_PIN_SCK          39
#define SD_MAX_TRANSFER_SZ  4000

esp_err_t sd_card_init(void) {
    // 1. 初始化 SPI 总线
    spi_bus_config_t bus_cfg = {
        .mosi_io_num = SD_PIN_MOSI,
        .miso_io_num = SD_PIN_MISO,
        .sclk_io_num = SD_PIN_SCK,
        .max_transfer_sz = SD_MAX_TRANSFER_SZ,
    };
    spi_bus_initialize(SD_HOST, &bus_cfg, SDSPI_DEFAULT_DMA);

    // 2. SDSPI 主机配置 (20MHz)
    sdmmc_host_t host = SDSPI_HOST_DEFAULT();
    host.slot = SD_HOST;
    host.max_freq_khz = 20000;

    // 3. 设备配置 (CS 引脚)
    sdspi_device_config_t slot_config = SDSPI_DEVICE_CONFIG_DEFAULT();
    slot_config.gpio_cs = SD_PIN_CS;

    // 4. 挂载 FATFS 文件系统
    esp_vfs_fat_sdmmc_mount_config_t mount_config = {
        .format_if_mount_failed = false,
        .max_files = 5,
        .allocation_unit_size = 16 * 1024,
    };
    return esp_vfs_fat_sdspi_mount(SD_MOUNT_POINT, &host,
        &slot_config, &mount_config, &s_card);
}'''

CODE_usb_descriptor = r'''// ===== usb_cdc.c — USB 复合设备描述符 (CDC + MSC) =====
static uint8_t const desc_configuration[] = {
    // 配置描述符: 1 配置, 3 接口, 远程唤醒, 100mA
    TUD_CONFIG_DESCRIPTOR(1, ITF_NUM_TOTAL, 0,
                          TUSB_DESC_TOTAL_LEN,
                          TUSB_DESC_CONFIG_ATT_REMOTE_WAKEUP, 100),
    // CDC 控制接口 + 数据接口 (虚拟串口)
    TUD_CDC_DESCRIPTOR(ITF_NUM_CDC, 4,
                       EDPT_CDC_NOTIF, 8,
                       EDPT_CDC_OUT, EDPT_CDC_IN, 64),
    // MSC 接口 (SD 卡读卡器)
    TUD_MSC_DESCRIPTOR(ITF_NUM_MSC, 5,
                       EDPT_MSC_OUT, EDPT_MSC_IN,
                       TUD_OPT_HIGH_SPEED ? 512 : 64),
};

// 设备描述符
static tusb_desc_device_t descriptor_config = {
    .bLength            = sizeof(descriptor_config),
    .bDescriptorType    = TUSB_DESC_DEVICE,
    .bcdUSB             = 0x0200,
    .bDeviceClass       = TUSB_CLASS_MISC,
    .bDeviceSubClass    = MISC_SUBCLASS_COMMON,
    .bDeviceProtocol    = MISC_PROTOCOL_IAD,
    .bMaxPacketSize0    = CFG_TUD_ENDPOINT0_SIZE,
    .idVendor           = 0x303A,   // Espressif VID
    .idProduct          = 0x4002,
    .bcdDevice          = 0x100,
    .iManufacturer      = 0x01,     // "Espressif"
    .iProduct           = 0x02,     // "ESP Recorder"
};

// USB 热插拔回调 — USB 插入后自动切换 SD 卡归属
void tud_mount_cb(void) {
    recorder_stop();         // 停止记录, 刷新缓冲
    sd_card_unmount_fs();    // 卸载 FATFS, SD 交给 USB MSC
}
void tud_umount_cb(void) {
    sd_card_remount_fs();    // USB 拔出, 重新挂载 FATFS
    if (cfg->recorder.auto_start) recorder_start();
}'''

CODE_wifi_ap = r'''// ===== wifi_config.c — AP 热点应用 =====
esp_err_t wifi_apply_ap(const char *ap_ssid, const char *ap_pass,
                        const char *ap_ip) {
    // 1. 设置静态 IP / 网关 / 掩码
    esp_netif_ip_info_t ip_info;
    ip_info.ip.addr      = ipaddr_addr(ap_ip ? ap_ip : "192.168.4.1");
    ip_info.gw.addr      = ipaddr_addr(ap_ip ? ap_ip : "192.168.4.1");
    ip_info.netmask.addr = ipaddr_addr("255.255.255.0");
    esp_netif_dhcps_stop(s_ap_netif);
    esp_netif_set_ip_info(s_ap_netif, &ip_info);
    esp_netif_dhcps_start(s_ap_netif);

    // 2. SSID: 入参优先, 否则自动生成 ESP_Recorder_XXXX
    char default_ssid[32];
    wifi_ap_default_ssid(default_ssid, sizeof(default_ssid));
    const char *ssid = (ap_ssid && ap_ssid[0]) ? ap_ssid : default_ssid;

    // 3. 构造 AP 配置 (WPA2-PSK, 信道 1, 最多 4 客户端)
    wifi_config_t ap_config = {
        .ap = { .ssid_len = 0, .channel = 1,
                .max_connection = 4,
                .authmode = is_open ? WIFI_AUTH_OPEN : WIFI_AUTH_WPA2_PSK },
    };
    strncpy((char*)ap_config.ap.ssid, ssid, 31);
    strncpy((char*)ap_config.ap.password, pass, 63);

    // 4. 应用配置 (AP+STA 混合模式)
    esp_wifi_set_mode(WIFI_MODE_APSTA);
    esp_wifi_set_config(WIFI_IF_AP, &ap_config);
    esp_wifi_start();
}'''

CODE_wifi_sta = r'''// ===== wifi_config.c — STA 连接外部 WiFi =====
esp_err_t wifi_connect_sta(const char *ssid, const char *pass,
                           int authmode) {
    wifi_config_t sta_config = {0};
    strncpy((char*)sta_config.sta.ssid, ssid, 31);
    strncpy((char*)sta_config.sta.password, pass, 63);
    sta_config.sta.threshold.authmode = (wifi_auth_mode_t)authmode;

    s_retry_count = 0;
    esp_wifi_disconnect();
    esp_wifi_set_config(WIFI_IF_STA, &sta_config);
    esp_wifi_start();
}
// WiFi 事件处理 — STA 断线自动重连 (最多 3 次, 间隔 5 秒)
static void wifi_event_handler(...) {
    if (event_id == WIFI_EVENT_STA_DISCONNECTED) {
        if (s_retry_count < 3) {
            s_retry_count++;
            vTaskDelay(pdMS_TO_TICKS(5000));
            esp_wifi_connect();  // 自动重连
        }
    }
}'''

CODE_recorder = r'''// ===== data_recorder.c — 双通道记录引擎核心 =====
// 每通道独立 16KB 环形缓冲 → 阈值/超时双条件落盘
#define SAVE_RINGBUF_SIZE   16384    // 16KB 环形缓冲
#define SAVE_THRESHOLD_BYTES 2048    // 满 2KB 立即刷盘
#define SAVE_TIMEOUT_MS     1000     // 超时 1 秒也刷盘

// RX 数据写入: 实时流 + 落盘环形缓冲
esp_err_t recorder_write_rx_channel(uint8_t channel,
                                    const uint8_t *data, size_t len) {
    recorder_live_push(channel, 0, data, len);  // 实时流环形缓冲
    recorder_write_raw(channel, data, len);     // 落盘环形缓冲
    return ESP_OK;
}
// 落盘任务 — 阈值/超时双条件触发 fwrite
static void recorder_io_task(void *arg) {
    while (1) {
        // 等待信号量 (有新数据 或 超时1秒)
        xSemaphoreTake(s_data_sem, pdMS_TO_TICKS(SAVE_TIMEOUT_MS));
        // 检查任一通道缓冲是否 >= 2048 字节
        for (uint8_t ch = 1; ch <= 2; ch++)
            if (ringbuf_count(&ch_state->save_rb) >= SAVE_THRESHOLD_BYTES)
                should_flush = true;
        if (should_flush) recorder_flush_all_locked(); // → fwrite(SD卡)
    }
}
// Sync 任务 — 每秒 fsync 保证数据安全
static void recorder_sync_task(void *arg) {
    while (1) {
        vTaskDelay(pdMS_TO_TICKS(1000));
        for (ch = 1; ch <= 2; ch++) {
            int rc = fsync(fileno(state->file));
            if (rc != 0 && ++state->fsync_fail_count >= 3) {
                state->failed = true;          // 3次失败 → 关闭文件
                recorder_close_channel_locked(ch);
}   }   }   }   }'''

CODE_recorder_file = r'''// ===== data_recorder.c — 文件名自动生成 =====
// 格式: /sdcard/REC_CH1_20260727_193428.bin
// 同名时自动加后缀: REC_CH1_20260727_193428_1.bin
static void recorder_compute_filename_locked(uint8_t channel) {
    struct timeval tv; gettimeofday(&tv, NULL);
    struct tm tm_info; localtime_r(&tv.tv_sec, &tm_info);
    for (int suffix = 0; suffix < 100; suffix++) {
        if (suffix == 0)
            snprintf(state->filename, sizeof(state->filename),
                "%s/REC_CH%u_%04d%02d%02d_%02d%02d%02d.bin",
                "/sdcard", channel,
                tm_info.tm_year+1900, tm_info.tm_mon+1, tm_info.tm_mday,
                tm_info.tm_hour, tm_info.tm_min, tm_info.tm_sec);
        else
            snprintf(state->filename, sizeof(state->filename),
                "%s/REC_CH%u_%04d%02d%02d_%02d%02d%02d_%d.bin", ...);
        if (stat(state->filename, &st) != 0) break; // 文件不存在, 可用
    }
}
// 实时流数据格式 (Web 端 200ms 拉取)
// 2026-07-27T19:34:28.123,CH1,RX,01 03 00 00 00 0A C5 CD
// 2026-07-27T19:34:28.234,CH2,TX,01 03 02 00 00 B8 44'''

CODE_router = r'''// ===== data_router.c — 跨通道双向转发 =====
// CH1 RX → CH2 TX   /   CH2 RX → CH1 TX
int data_router_forward(uint8_t source_channel,
                        const uint8_t *data, size_t len) {
    // 确定目标通道
    uint8_t target = (source_channel == USER_UART_CHANNEL_1)
        ? USER_UART_CHANNEL_2 : USER_UART_CHANNEL_1;

    // 通过 UART 发送到目标通道
    int sent = user_uart_write_channel(target, data, len);

    // 更新统计 (线程安全)
    portENTER_CRITICAL(&s_stats_lock);
    if (source_channel == USER_UART_CHANNEL_1) {
        s_stats.ch1_to_ch2_bytes += sent;
        s_stats.ch1_to_ch2_drop_bytes += (len - sent);
    } else {
        s_stats.ch2_to_ch1_bytes += sent;
        s_stats.ch2_to_ch1_drop_bytes += (len - sent);
    }
    portEXIT_CRITICAL(&s_stats_lock);
    return sent;
}'''

CODE_web_handlers = r'''// ===== web_server.c — HTTP API 端点示例 (14 个) =====

// ① 设备状态 (JSON 返回全部运行指标)
static esp_err_t status_handler(httpd_req_t *req) {
    cJSON *root = cJSON_CreateObject();
    cJSON_AddBoolToObject(root, "ap", ws->ap_started);
    cJSON_AddBoolToObject(root, "sta_connected", ws->sta_connected);
    cJSON_AddBoolToObject(root, "sd", sd->mounted);
    cJSON_AddBoolToObject(root, "recording", recorder_is_running());
    cJSON_AddNumberToObject(root, "rec_total_bytes",
                           (double)recorder_get_total_bytes());
    // ... 共计 25+ 个状态字段
    char *body = cJSON_PrintUnformatted(root);
    httpd_resp_set_type(req, "application/json; charset=utf-8");
    return httpd_resp_send(req, body, strlen(body));
}

// ② 设置 STA WiFi → 写入 config.json → 触发重连
static esp_err_t wifi_set_handler(httpd_req_t *req) {
    // 解析 JSON body: {"ssid":"...", "pass":"...", "auth":3}
    cJSON *root = cJSON_Parse(buf);
    const char *ssid = cJSON_GetObjectItem(root,"ssid")->valuestring;
    // 只覆盖 STA 字段, 其他保留
    config_manager_load(mgr, "/sdcard/config.json");
    strncpy(mgr->cfg.wifi.sta_ssid, ssid, 31);
    strncpy(mgr->cfg.wifi.sta_pass, pass, 63);
    config_manager_save(mgr, "/sdcard/config.json");
    wifi_connect_sta(ssid, pass, authmode);  // 触发连接
    return httpd_resp_send(req, "{\"ok\":true}", 10);
}

// ③ HEX 发送到指定通道 TX
static esp_err_t send_tx_handler(httpd_req_t *req) {
    // POST body: {"channel":1, "hex":"01 03 00 00 00 0A"}
    uint8_t buf[256]; size_t out = 0;
    for (const char *p = hex; *p && out < sizeof(buf); p++) {
        // 解析 HEX 字符串 → 二进制字节
        int hi = hex_nibble(*p), lo = hex_nibble(*(p+1));
        buf[out++] = (uint8_t)((hi << 4) | lo); p++;
    }
    int sent = user_uart_write_channel(channel, buf, out);
    // 返回: {"ok":true, "channel":1, "sent":6}
}}

// ④ 实时数据流增量拉取 (200ms 轮询)
static esp_err_t stream_fetch_handler(httpd_req_t *req) {
    uint32_t since = get_query_value(req, "since"); // ?since=123
    recorder_live_item_t items[32];
    size_t got = recorder_live_fetch(since, items, 32, &max_seq);
    // 返回: {"items":[{"seq":1,"line":"...RX,01 03..."}], "max_seq":5}
}

// ⑤ 文件下载 (触发浏览器另存为)
static esp_err_t file_download_handler(httpd_req_t *req) {
    char path[128]; snprintf(path, sizeof(path), "/sdcard/%s", name);
    FILE *f = fopen(path, "r");
    httpd_resp_set_type(req, "application/octet-stream");
    httpd_resp_set_hdr(req, "Content-Disposition",
                       "attachment; filename=\"...\"");
    while ((n = fread(chunk, 1, 1024, f)) > 0)
        httpd_resp_send_chunk(req, chunk, n);  // 分块传输
}'''

CODE_web_routes = r'''// ===== web_server.c — 14 个 API 路由注册 =====
static const httpd_uri_t uris[] = {
    {.uri="/",                   .method=HTTP_GET,  .handler=index_handler},
    {.uri="/api/status",         .method=HTTP_GET,  .handler=status_handler},
    {.uri="/api/wifi",           .method=HTTP_POST, .handler=wifi_set_handler},
    {.uri="/api/ap",             .method=HTTP_POST, .handler=ap_set_handler},
    {.uri="/api/files",          .method=HTTP_GET,  .handler=files_list_handler},
    {.uri="/api/file",           .method=HTTP_GET,  .handler=file_download_handler},
    {.uri="/api/file/delete",    .method=HTTP_POST, .handler=file_delete_handler},
    {.uri="/api/recorder/start", .method=HTTP_POST, .handler=recorder_start_handler},
    {.uri="/api/recorder/stop",  .method=HTTP_POST, .handler=recorder_stop_handler},
    {.uri="/api/stream",         .method=HTTP_GET,  .handler=stream_fetch_handler},
    {.uri="/api/stream/clear",   .method=HTTP_POST, .handler=stream_clear_handler},
    {.uri="/api/send",           .method=HTTP_POST, .handler=send_tx_handler},
    {.uri="/api/reboot",         .method=HTTP_POST, .handler=reboot_handler},
    {.uri="/api/config/stream_gap", .method=HTTP_POST, .handler=stream_gap_set_handler},
};
for (size_t i = 0; i < sizeof(uris)/sizeof(uris[0]); i++)
    httpd_register_uri_handler(s_server, &uris[i]);'''

CODE_config = r'''// ===== config_manager.c — JSON 配置加载 =====
esp_err_t config_manager_load(config_manager_t* manager, const char* path) {
    // 1. 先填充默认值
    config_manager_reset_defaults(manager);
    // 2. 读取 /sdcard/config.json
    FILE* f = fopen(path, "r");
    if (f == NULL) {
        manager->loaded = false;
        return ESP_ERR_NOT_FOUND;  // 不存在 → 用默认值
    }
    // 3. 读取全部内容
    char* buf = malloc(size + 1);
    fread(buf, 1, size, f); fclose(f);
    // 4. 解析 JSON
    cJSON* root = cJSON_Parse(buf);
    // 5. 逐字段覆盖 (缺失字段保留默认值)
    cJSON* uart = cJSON_GetObjectItem(root, "uart");
    if (uart) {
        manager->cfg.uart.baudrate = get_json_int(uart, "baudrate", 115200);
        // ... databits, stopbits, parity
    }
    cJSON* wifi = cJSON_GetObjectItem(root, "wifi");
    if (wifi) {
        copy_string(manager->cfg.wifi.sta_ssid, 32,
                    get_json_string(wifi, "sta_ssid", ""));
        manager->cfg.wifi.sta_authmode = clamp_authmode(
            get_json_int(wifi, "sta_authmode", 0));
        // ... ap_ssid, ap_pass, ap_ip
    }
    cJSON* recorder = cJSON_GetObjectItem(root, "recorder");
    if (recorder) {
        manager->cfg.recorder.auto_start =
            get_json_int(recorder, "auto_start", 1) != 0;
        manager->cfg.recorder.stream_gap_ms =
            config_manager_clamp_stream_gap_ms(
                get_json_int(recorder, "stream_gap_ms", 100));
    }
    manager->loaded = true;
}'''

CODE_console = r'''// ===== console_cmd.c — 调试控制台 (UART1 115200) =====
esp_err_t console_start(void) {
    esp_console_repl_config_t repl_cfg = ESP_CONSOLE_REPL_CONFIG_DEFAULT();
    repl_cfg.prompt = "esp> ";
    esp_console_dev_uart_config_t hw = ESP_CONSOLE_DEV_UART_CONFIG_DEFAULT();
    esp_console_new_repl_uart(&hw, &repl_cfg, &repl);
    register_sd_test();          // 注册 sd_test 命令
    esp_console_start_repl(repl);
}
// sd_test 命令 — SD 卡读写速度测试
static int cmd_sd_test(int argc, char **argv) {
    if (argc == 1)
        return sd_speed_test_run_all();       // 标准 1/4/16/64 KB
    if (argc == 3)
        return sd_speed_test_once(block, total);  // 自定义参数
}
// 使用示例:
// esp> sd_test
// esp> sd_test 4096 1048576'''

CODE_html_js = r'''// ===== index.html — Web 前端关键 JavaScript =====

// 实时流 200ms 增量轮询
async function pollStream() {
    const r = await fetch('/api/stream?since=' + sinceSeq);
    const j = await r.json();
    for (const it of j.items) {
        // 超过 stream_gap_ms 未收到数据 → 插入 "等待数据..."
        if (lastItemTs && (it.ts - lastItemTs) > streamGapMs) {
            el.appendChild(makeGap('等待数据... ('+gap+'ms)'));
        }
        el.appendChild(makeLine(it.line));  // 显示 HEX 行
        sinceSeq = Math.max(sinceSeq, it.seq);
    }
    el.scrollTop = el.scrollHeight;  // 自动滚动到底部
}
setInterval(pollStream, 200);

// 发送 HEX 到指定通道
async function sendTx() {
    const hex = document.getElementById('txdata').value.trim();
    const channel = Number(document.getElementById('txchannel').value);
    const r = await api('/api/send', {method:'POST',
        body: JSON.stringify({channel, hex})});
}

// 状态自动刷新 (2 秒间隔)
async function refreshStatus() {
    const s = await api('/api/status');
    document.getElementById('ap').innerHTML =
        s.ap ? '<span class=on>已开启</span>' : '<span class=off>未开启</span>';
    document.getElementById('sd').innerHTML =
        s.sd ? '<span class=on>已挂载</span>' : '<span class=off>未挂载</span>';
    document.getElementById('rec').innerHTML =
        s.recording ? '<span class=on>记录中</span>' : '<span class=off>空闲</span>';
    document.getElementById('ch1rx').textContent = s.uart_ch1_rx_bytes || 0;
    document.getElementById('ch2rx').textContent = s.uart_ch2_rx_bytes || 0;
    // ... 30+ 个状态字段更新
}
setInterval(refreshStatus, 2000);

// 自动限制显示行数 (最多 500 行)
function trimLines() {
    if (liveLines.length <= 500) return;
    liveLines = liveLines.slice(-500);
    document.getElementById('stream').replaceChildren(
        ...liveLines.map(x => x.node));
}'''

CODE_cmake = r'''# ===== main/CMakeLists.txt — 编译源文件清单 =====
idf_component_register(
    SRCS "app_main.c" "console_cmd.c" "sd_speed_test.c"
         "web_server.c" "config_manager.c" "data_recorder.c"
         "data_router.c" "wifi_config.c" "usb_cdc.c"
         "usb_msc.c" "uart_driver.c" "sd_card.c"
    INCLUDE_DIRS "."
    EMBED_FILES "index.html")   # ← 内嵌到固件

# ===== 顶层 CMakeLists.txt =====
cmake_minimum_required(VERSION 3.16)
include($ENV{IDF_PATH}/tools/cmake/project.cmake)
project(esp_recorder_v1.0)'''

# ====================== Word 辅助函数 ======================

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_style(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]; run = p.add_run(header)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1: set_cell_shading(cell, 'D6E4F0')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table

def add_code_block(doc, code_text, title=None):
    """添加代码块"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(code_text.strip().split('\n')):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        pPr = p._p.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        pPr.append(shading)
        run = p.add_run(line)
        run.font.name = 'Consolas'; run.font.size = Pt(7.5)
    doc.add_paragraph()

def add_photo_placeholder(doc, caption, width_cm=14):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(width_cm)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '3600')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('📷 请在此处插入照片'); run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88); run.italic = True
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'【{caption}】'); run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x2F, 0x54, 0x96); run2.bold = True
    doc.add_paragraph()

def add_heading_styled(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_normal(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    run.font.size = Pt(10.5)
    return p

def add_page_break(doc):
    doc.add_page_break()


# ====================== 主构建函数 ======================

def build_docx():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)

    # ============ 封面 ============
    for _ in range(6): doc.add_paragraph()
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ESP Recorder'); run.bold = True
    run.font.size = Pt(32); run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('ESP32-S3 双路串口数据记录仪\n项目开发完整报告')
    run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph(); doc.add_paragraph()
    for label, value in [
        ('项目名称', 'esp_recorder_v1.0'), ('主控芯片', 'ESP32-S3'),
        ('开发框架', 'ESP-IDF v5.4'),
        ('开发工具', 'VS Code + Espressif IDF Extension v2.1.0'),
        ('文档日期', '2026-08-03'), ('作者', '尤译庆')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}'); run.font.size = Pt(11)
    add_page_break(doc)

    # ============ 目录 ============
    add_heading_styled(doc, '目录', level=1)
    for item in [
        '1. 项目概述 (定位/场景/结构/入口代码)',
        '2. 功能清单 (16项功能表格)',
        '3. 硬件接线 (SD卡/串口/控制台/USB 引脚表)',
        '4. 软件架构 (初始化流程/FreeRTOS任务/数据流)',
        '5. 环境搭建 (ESP-IDF安装/VS Code配置/esp_idf.json)',
        '6. 编译构建与烧录 (命令/产物/Flash布局/代码片段)',
        '7. 配置文件说明 (config.json 完整示例)',
        '8. 各模块详细代码 (10个模块源码+解析)',
        '9. Web 管理界面 (API路由/前端JS代码/截图区)',
        '10. 数据流与工作流程',
        '11. 调试与常见问题',
        '12. 已知问题与注意事项',
        '附录: VS Code 环境恢复步骤 / CMakeLists.txt',
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].font.size = Pt(11)
    add_page_break(doc)

    # ============ 第1章 项目概述 ============
    add_heading_styled(doc, '1. 项目概述', level=1)

    add_heading_styled(doc, '1.1 项目定位', level=2)
    add_normal(doc, 'ESP Recorder 是一个基于 ESP32-S3 的双路串口数据记录仪。它能够同时监听两路独立 UART 串口的数据，将所有接收到的原始字节实时保存到 SD 卡，并通过 WiFi 热点提供 Web 管理界面，支持远程查看状态、启停记录、下载文件和实时监控数据流。')

    add_heading_styled(doc, '1.2 核心应用场景', level=2)
    for s in ['工业设备串口通信数据采集与记录', '嵌入式系统调试时的长时间数据抓取',
              '多路传感器数据的集中采集存储', '需要"黑匣子"功能的串口通信场景']:
        add_bullet(doc, s)

    add_photo_placeholder(doc, '图1-1: 项目整体效果或硬件实物照片')

    add_heading_styled(doc, '1.3 项目目录结构', level=2)
    add_table_with_style(doc, ['文件/目录', '说明'], [
        ['esp_recorder_v1.1/CMakeLists.txt', '顶层 CMake (项目入口)'],
        ['esp_recorder_v1.1/sdkconfig', 'ESP-IDF 配置 (esp32s3, 各外设参数)'],
        ['main/app_main.c', '★ 主入口 — 13步初始化所有模块'],
        ['main/index.html', '★ Web 前端单页应用 (EMBED_FILES 内嵌固件)'],
        ['main/uart_driver.c/h', '双路 UART 驱动 (CH1/CH2 独立初始化)'],
        ['main/sd_card.c/h', 'SD 卡 SPI 驱动 + FATFS 挂载/卸载/重挂载'],
        ['main/usb_cdc.c/h', 'USB CDC 虚拟串口 (TinyUSB 复合设备)'],
        ['main/usb_msc.c/h', 'USB MSC 读卡器 (SCSI READ10/WRITE10)'],
        ['main/wifi_config.c/h', 'WiFi AP+STA 混合模式管理'],
        ['main/web_server.c/h', 'HTTP REST API 服务 (端口80, 14个端点)'],
        ['main/data_recorder.c/h', '双通道数据记录引擎 (环形缓冲→落盘)'],
        ['main/data_router.c/h', '跨通道数据路由 (CH1↔CH2 双向转发)'],
        ['main/config_manager.c/h', 'JSON 配置文件读写 (/sdcard/config.json)'],
        ['main/console_cmd.c/h', '串口调试控制台 (sd_test 命令)'],
        ['main/sd_speed_test.c/h', 'SD 卡读写速度测试'],
        ['components/', 'TinyUSB 协议栈 (Espressif 封装)'],
        ['artifacts/', '构建日志归档 (6个历史构建日志)'],
        ['.vscode/', 'VS Code 工作区配置 (settings/launch/c_cpp)'],
    ], col_widths=[6, 10])

    add_heading_styled(doc, '1.4 主入口核心代码 (app_main.c)', level=2)
    add_normal(doc, '以下是项目主入口的完整初始化流程代码，展示了所有子系统的启动顺序：')
    add_code_block(doc, CODE_app_main)
    add_code_block(doc, CODE_full_init_flow)

    # ============ 第2章 功能清单 ============
    add_heading_styled(doc, '2. 功能清单 (16 项)', level=1)
    add_table_with_style(doc, ['编号', '功能模块', '描述', '状态'], [
        ['F01', '双路 UART 采集', 'CH1(TX43/RX44)+CH2(TX17/RX18),独立初始化,一通道失败不影响另一通道', '✅ 编译通过'],
        ['F02', 'SD卡 FATFS存储', 'SPI模式(CS10/MOSI38/MISO40/SCK39),REC_CH*.bin自动命名', '✅ 编译通过'],
        ['F03', 'WiFi AP热点', 'SSID=ESP_Recorder_XXXX, 密码12345678, IP=192.168.4.1,最多4客户端', '✅ 编译通过'],
        ['F04', 'WiFi STA连接', 'WPA2/WPA3/开放, 自动重连3次间隔5秒, 断线不影响AP', '✅ 编译通过'],
        ['F05', 'Web 管理界面', '内嵌index.html单页应用,端口80,零外部依赖', '✅ 编译通过'],
        ['F06', 'REST API', '14个端点(状态/配网/AP配置/文件列表/下载/删除/记录启停/实时流/发送TX/重启等)', '✅ 编译通过'],
        ['F07', '实时数据流', '256条环形缓冲,200ms增量拉取,HEX格式,RX/TX方向标记', '✅ 编译通过'],
        ['F08', '数据路由', 'CH1 RX→CH2 TX, CH2 RX→CH1 TX, 硬件级双向转发,独立丢包统计', '✅ 编译通过'],
        ['F09', 'USB 复合设备', 'TinyUSB CDC(虚拟串口)+MSC(SD读卡器),VID=303A PID=4002', '✅ 编译通过'],
        ['F10', 'USB 热插拔保护', 'USB插入→自动卸载FATFS停止记录, 拔出→重挂载FATFS恢复记录', '✅ 编译通过'],
        ['F11', 'JSON 配置系统', '/sdcard/config.json持久化,含UART/WiFi/Recorder三段,缺失字段默认值', '✅ 编译通过'],
        ['F12', '记录启停控制', 'Web API远程控制, auto_start上电自启动, 阈值/超时双条件落盘', '✅ 编译通过'],
        ['F13', '双通道独立统计', '每通道独立RX/TX字节数、丢字节数、文件打开状态、落盘文件名', '✅ 编译通过'],
        ['F14', '调试控制台', 'UART1(TX1/RX2) 115200baud, sd_test命令测试SD读写速度', '✅ 编译通过'],
        ['F15', '文件管理', 'Web端列出REC_*.bin/下载/删除, 路径穿越防护', '✅ 编译通过'],
        ['F16', '远程重启', 'Web API→esp_restart(), 200ms延迟保证响应发出', '✅ 编译通过'],
    ], col_widths=[1.0, 3.0, 8.5, 3.0])

    # ============ 第3章 硬件接线 ============
    add_heading_styled(doc, '3. 硬件接线', level=1)

    add_photo_placeholder(doc, '图3-1: ESP32-S3开发板与SD卡模块、串口设备的完整接线照片')
    add_photo_placeholder(doc, '图3-2: 接线细节特写 (SD卡SPI引脚)')

    add_heading_styled(doc, '3.1 SD 卡 (SPI 模式)', level=2)
    add_table_with_style(doc, ['SD卡引脚', 'ESP32-S3 GPIO', '说明'], [
        ['CS', 'GPIO 10', '片选信号'], ['MOSI', 'GPIO 38', '主机输出→从机输入'],
        ['MISO', 'GPIO 40', '从机输出→主机输入'], ['SCK', 'GPIO 39', '时钟信号'],
        ['VCC', '3.3V', '供电 (勿接5V)'], ['GND', 'GND', '共地'],
    ], col_widths=[4, 5, 7])
    add_normal(doc, 'SPI主机: SPI2_HOST, 最大频率 20MHz, 最大单次传输 4000 字节。')

    add_heading_styled(doc, '3.2 用户串口', level=2)
    add_table_with_style(doc, ['通道', 'UART编号', 'TX Pin', 'RX Pin', '默认波特率', '缓冲大小'], [
        ['CH1', 'UART0', 'GPIO 43', 'GPIO 44', '115200', '2048 B'],
        ['CH2', 'UART2', 'GPIO 17', 'GPIO 18', '115200', '2048 B'],
    ], col_widths=[2, 2.5, 3, 3, 3, 2.5])

    add_heading_styled(doc, '3.3 引脚定义代码 (uart_driver.h)', level=2)
    add_normal(doc, '以下是项目中 uart_driver.h 的实际引脚宏定义：')
    add_code_block(doc, CODE_uart_h)

    add_heading_styled(doc, '3.4 其他 GPIO', level=2)
    add_table_with_style(doc, ['功能', 'GPIO', '方向', '说明'], [
        ['UART接收使能', 'GPIO 16', '输出(低)', '上电拉低使能串口接收'],
        ['USB VBUS检测', 'GPIO 48', '输入', 'TinyUSB VBUS Monitor (检测USB连接)'],
        ['调试控制台TX', 'GPIO 1', '输出', 'UART1 控制台输出 (115200)'],
        ['调试控制台RX', 'GPIO 2', '输入', 'UART1 控制台输入 (115200)'],
    ], col_widths=[4, 3, 3, 6])

    add_heading_styled(doc, '3.5 USB 接口', level=2)
    add_normal(doc, 'ESP32-S3 原生 USB OTG 接口，枚举为复合设备：')
    add_bullet(doc, 'CDC (虚拟串口): "ESP Recorder CDC" — 可用于调试或配网备用通道')
    add_bullet(doc, 'MSC (大容量存储): "ESP Recorder MSC" — SD卡作为U盘暴露给电脑')
    add_bullet(doc, 'VID: 303A (Espressif 官方VID), PID: 4002')
    add_normal(doc, '⚠️ 使用开发板的 USB OTG 口(非 UART 口)连接电脑才能枚举 USB 复合设备。如果使用 USB-TTL 模块，需交叉 TX/RX、共地、3.3V电平、不要给VCC供电。')

    # ============ 第4章 软件架构 ============
    add_heading_styled(doc, '4. 软件架构', level=1)

    add_heading_styled(doc, '4.1 完整初始化流程', level=2)
    add_code_block(doc, CODE_full_init_flow)

    add_heading_styled(doc, '4.2 FreeRTOS 任务一览', level=2)
    add_table_with_style(doc, ['任务名称', '优先级', '栈大小', '核心功能'], [
        ['uart_rx_ch1', '5', '4096', 'CH1 UART接收→recorder落盘+router转发到CH2'],
        ['uart_rx_ch2', '5', '4096', 'CH2 UART接收→recorder落盘+router转发到CH1'],
        ['rec_io', '5', '4096', '环形缓冲→SD卡fwrite (阈值2KB/超时1秒双条件触发)'],
        ['rec_sync', '1', '3072', '每秒fsync强制同步, 连续3次失败→关闭文件标记故障'],
        ['tinyusb', '5', '4096', 'TinyUSB协议栈 (CPU1亲和, 处理CDC+MSC复合设备)'],
    ], col_widths=[3.5, 2, 2, 8.5])

    add_heading_styled(doc, '4.3 数据流架构图', level=2)
    add_normal(doc, '完整数据流路径 (从外部设备到 SD 卡 + Web 显示 + 跨通道转发)：')
    data_flow = (
        '外部设备TX ──→ CH1 UART(RX44) ──→ uart_rx_ch1 task\n'
        '    ├── recorder_write_rx_channel()\n'
        '    │     ├── live_push() → 实时流环形缓冲(256条) → Web /api/stream 拉取\n'
        '    │     └── ringbuf_write() → 落盘环形缓冲(16KB)\n'
        '    │            └── rec_io task → 阈值/超时触发 → fwrite(SD卡 REC_CH1_*.bin)\n'
        '    └── data_router_forward() → CH2 UART(TX17) → 另一设备\n'
        '\n'
        '外部设备TX ──→ CH2 UART(RX18) ──→ uart_rx_ch2 task\n'
        '    ├── recorder_write_rx_channel() → live_push() + ringbuf_write() → SD卡\n'
        '    └── data_router_forward() → CH1 UART(TX43) → 另一设备\n'
        '\n'
        'Web /api/send POST → user_uart_write_channel() → TX发送 + recorder_write_tx_channel()\n'
        'Web /api/stream GET  → recorder_live_fetch(since) → JSON {items, max_seq}\n'
        '\n'
        'USB插入: tud_mount_cb() → recorder_stop() → sd_card_unmount_fs() → SD交MSC\n'
        'USB拔出: tud_umount_cb() → sd_card_remount_fs() → recorder_start()'
    )
    add_code_block(doc, data_flow)

    add_heading_styled(doc, '4.4 UART RX 任务代码', level=2)
    add_normal(doc, '每个 UART 通道的接收任务逻辑完全相同，收到数据后同时分发到记录引擎和路由器：')
    add_code_block(doc, CODE_uart_rx_task)

    # ============ 第5章 环境搭建 ============
    add_heading_styled(doc, '5. 环境搭建步骤', level=1)

    add_heading_styled(doc, '5.1 安装 ESP-IDF v5.4', level=2)
    add_bullet(doc, '推荐使用 VS Code Espressif IDF 扩展的安装管理器自动安装')
    add_bullet(doc, '或从 Espressif 官网 (https://docs.espressif.com/projects/esp-idf/) 下载离线安装包')
    add_bullet(doc, '本项目实际安装路径: D:\\23178\\esp-idf')
    add_bullet(doc, 'Python 虚拟环境: C:\\Users\\23178\\.espressif\\python_env\\idf5.4_py3.8_env')

    add_heading_styled(doc, '5.2 安装 VS Code 扩展', level=2)
    add_bullet(doc, '在 VS Code 扩展商店搜索 "Espressif IDF"')
    add_bullet(doc, '本项目使用版本: espressif.esp-idf-extension@2.1.0')
    add_bullet(doc, '扩展会自动检测 ESP-IDF 安装并配置工具链')

    add_heading_styled(doc, '5.3 验证安装命令', level=2)
    add_code_block(doc, 'idf.py --version\npython --version\ncode --list-extensions --show-versions | Select-String espressif')

    add_heading_styled(doc, '5.4 VS Code 工作区配置', level=2)
    add_normal(doc, '.vscode/settings.json 的关键配置内容：')
    add_code_block(doc, '''{
    "idf.currentSetup": "D:\\\\23178\\\\esp-idf",
    "idf.customExtraVars": { "IDF_TARGET": "esp32s3" },
    "idf.flashType": "UART"
}''')
    add_normal(doc, '⚠️ 重要: 不要在工作区预设 idf.portWin，让每次烧录时根据硬件枚举结果选择正确的 COM 口。避免重复之前 COM90(蓝牙虚拟串口) 的错误。')

    add_heading_styled(doc, '5.5 ESP-IDF 安装注册表', level=2)
    add_normal(doc, 'VS Code 扩展通过以下文件识别已安装的工具链：')
    add_code_block(doc, '''C:\\Users\\23178\\.espressif\\esp_idf.json
{
    "idfSelectedId": "esp-idf-v5.4",
    "setups": [{
        "id": "esp-idf-v5.4",
        "idfPath": "D:\\\\23178\\\\esp-idf",
        "python": "...\\python_env\\idf5.4_py3.8_env\\Scripts\\python.exe"
    }]
}''')
    add_photo_placeholder(doc, '图5-1: VS Code 打开项目后的界面截图')
    add_photo_placeholder(doc, '图5-2: ESP-IDF 扩展配置界面截图')

    # ============ 第6章 编译构建与烧录 ============
    add_heading_styled(doc, '6. 编译构建与烧录', level=1)

    add_heading_styled(doc, '6.1 构建命令', level=2)
    add_code_block(doc, '''cd D:\\zhuomian\\weite\\ESP32\\esp_recorder_v1.1
idf.py build
# 保存完整构建日志
idf.py build 2>&1 | Tee-Object -FilePath artifacts\\idf_build.txt
# 清理后重新构建
idf.py fullclean && idf.py build''')

    add_heading_styled(doc, '6.2 编译产物', level=2)
    add_table_with_style(doc, ['文件', '说明', '大小'], [
        ['build/esp_recorder_v1.0.bin', '主应用程序固件', '~0xFA000 (约1MB)'],
        ['build/bootloader/bootloader.bin', '启动引导程序', '~16KB'],
        ['build/partition_table/partition-table.bin', '分区表', '3KB'],
        ['build/ota_data_initial.bin', 'OTA初始数据', '8KB'],
        ['build/flash_args', '烧录参数清单 (含偏移地址)', '—'],
        ['build/*.elf', 'ELF调试文件 (含符号表)', '—'],
        ['build/*.map', '链接映射文件', '—'],
    ], col_widths=[6.5, 6.5, 3])

    add_heading_styled(doc, '6.3 Flash 分区布局', level=2)
    add_table_with_style(doc, ['偏移地址', '分区名称', '内容', '大小'], [
        ['0x0', 'bootloader', 'bootloader/bootloader.bin', '32KB'],
        ['0x8000', 'partition_table', 'partition-table.bin', '4KB'],
        ['0xf000', 'ota_data', 'ota_data_initial.bin', '8KB'],
        ['0x20000', 'factory', 'esp_recorder_v1.0.bin', '4MB'],
    ], col_widths=[3, 4, 6, 3])
    add_normal(doc, 'Flash参数: dio模式, 80MHz频率, 16MB Flash容量, 编译后固件76%空间空闲。')

    add_heading_styled(doc, '6.4 烧录命令', level=2)
    add_code_block(doc, '''# 仅烧录 (不擦除已有数据)
idf.py -p COMx flash
# 烧录并打开串口监视器 (查看日志输出)
idf.py -p COMx flash monitor
# 完全擦除后烧录 (⚠️ 谨慎使用)
idf.py -p COMx erase-flash flash''')
    add_normal(doc, '⚠️ 烧录前务必确认: (1) COM口是ESP32-S3真实串口,不是蓝牙虚拟串口 (2) 芯片型号是ESP32-S3 (3) Flash容量16MB匹配。')

    add_heading_styled(doc, '6.5 CMakeLists.txt 编译配置', level=2)
    add_normal(doc, '以下展示项目两个 CMakeLists.txt 的完整内容：')
    add_code_block(doc, CODE_cmake)

    add_photo_placeholder(doc, '图6-1: VS Code终端 idf.py build 构建成功截图 (显示 Project build complete)')
    add_photo_placeholder(doc, '图6-2: 烧录过程的终端截图')

    # ============ 第7章 配置文件 ============
    add_heading_styled(doc, '7. 配置文件说明', level=1)

    add_normal(doc, '配置文件路径: /sdcard/config.json (SD 卡根目录)。首次开机如 SD 卡上无此文件，系统自动创建默认配置。')
    add_heading_styled(doc, '7.1 完整 JSON 配置', level=2)
    add_code_block(doc, '''{
  "uart": {
    "baudrate": 115200,        // 两路UART共用波特率 (8N1固定格式)
    "databits": 8,             // 数据位 (当前固定8)
    "stopbits": 1,             // 停止位 (当前固定1)
    "parity": 0                // 0=无校验, 1=奇校验, 2=偶校验
  },
  "wifi": {
    "enable_ap": 1,            // 是否启用AP热点
    "ap_ssid": "",             // AP名称 (空=自动生成 ESP_Recorder_XXXX)
    "ap_pass": "12345678",     // AP密码 (空=开放热点,无需密码)
    "ap_ip": "192.168.4.1",    // AP静态IP
    "enable_sta": 0,           // 是否启用STA连接外部WiFi
    "sta_ssid": "",            // 外部WiFi SSID
    "sta_pass": "",            // 外部WiFi密码
    "sta_authmode": 0          // 加密: 0=OPEN, 3=WPA2-PSK, 6=WPA3, 7=WPA2/WPA3过渡
  },
  "recorder": {
    "auto_start": 1,           // 上电自动开始记录
    "stream_gap_ms": 100       // 实时流间隔阈值(ms), 超出显示"等待数据..."
  }
}''')

    add_heading_styled(doc, '7.2 配置加载代码 (config_manager.c)', level=2)
    add_normal(doc, '以下是配置管理器的核心加载逻辑，展示了 JSON 解析和字段默认值机制：')
    add_code_block(doc, CODE_config)
    add_photo_placeholder(doc, '图7-1: SD卡中 config.json 文件实际内容截图')

    # ============ 第8章 各模块代码详解 ============
    add_heading_styled(doc, '8. 各模块详细代码', level=1)
    add_normal(doc, '本章逐一展示项目 10 个核心模块的实际代码片段和设计说明。', bold=True)

    # 8.1
    add_heading_styled(doc, '8.1 app_main.c — 主入口 (13步初始化)', level=2)
    add_normal(doc, '主入口是整个系统的启动核心，按严格顺序初始化13个子系统。每个步骤的失败处理独立。')
    add_code_block(doc, CODE_app_main, '▼ 主入口完整代码')

    # 8.2
    add_heading_styled(doc, '8.2 uart_driver.c — 双路 UART 驱动', level=2)
    add_normal(doc, '两路 UART 完全独立，支持一通道初始化失败时另一通道不受影响，运行时动态修改波特率。使用 portMUX_TYPE 自旋锁保证多任务间的 RX/TX 计数器线程安全。')
    add_code_block(doc, CODE_uart_h, '▼ 引脚定义与接口 (uart_driver.h)')
    add_code_block(doc, CODE_uart_init, '▼ 单通道初始化 (uart_driver.c)')

    # 8.3
    add_heading_styled(doc, '8.3 sd_card.c — SD 卡 SPI 驱动', level=2)
    add_normal(doc, 'SPI 模式访问 SD 卡，FATFS 挂载到 /sdcard。关键特性是支持安全卸载/重挂载: USB MSC 接管时卸载 FATFS 但保留 SPI 块设备；USB 断开后重挂载并恢复记录。')
    add_code_block(doc, CODE_sd_spi, '▼ SD卡 SPI 初始化完整代码')

    # 8.4
    add_heading_styled(doc, '8.4 usb_cdc.c + usb_msc.c — USB 复合设备', level=2)
    add_normal(doc, 'TinyUSB 实现 CDC (虚拟串口) + MSC (SD读卡器) 复合设备，一根 USB 线同时提供调试通道和文件传输。实现完整 SCSI READ10/WRITE10 命令，3次重试机制。')
    add_code_block(doc, CODE_usb_descriptor, '▼ USB 复合设备描述符与热插拔回调')

    # 8.5
    add_heading_styled(doc, '8.5 wifi_config.c — WiFi AP+STA 管理', level=2)
    add_normal(doc, 'AP+STA 混合模式 (WIFI_MODE_APSTA)，AP 永远可用保证本地 Web 管理，STA 可选连接外部网络。STA 断线自动重连最多3次间隔5秒。')
    add_code_block(doc, CODE_wifi_ap, '▼ AP 热点启动代码')
    add_code_block(doc, CODE_wifi_sta, '▼ STA 连接与自动重连代码')

    # 8.6
    add_heading_styled(doc, '8.6 data_recorder.c — 数据记录引擎 (最核心模块)', level=2)
    add_normal(doc, '这是项目最核心的模块。双通道独立记录，每通道 16KB 环形缓冲 → 双条件落盘 (满2KB或超时1秒)。每秒 fsync 保证数据安全。实时流 256 条环形缓冲供 Web 端增量拉取。')
    add_code_block(doc, CODE_recorder, '▼ 记录引擎核心代码 (环形缓冲 + 双条件落盘 + fsync)')
    add_code_block(doc, CODE_recorder_file, '▼ 文件名自动生成 + 实时流数据格式')

    # 8.7
    add_heading_styled(doc, '8.7 data_router.c — 跨通道数据路由', level=2)
    add_normal(doc, '两个外部设备分别接 CH1 和 CH2，路由器自动将 CH1 收到的数据转发到 CH2 TX，CH2 收到的数据转发到 CH1 TX，实现透明双向通信，同时全程记录。')
    add_code_block(doc, CODE_router, '▼ 跨通道路由转发代码')

    # 8.8
    add_heading_styled(doc, '8.8 web_server.c — HTTP 服务器 (14个API)', level=2)
    add_normal(doc, '基于 esp_http_server 组件，端口 80，14 个 REST API 端点。使用 cJSON 安全拼装 JSON（避免 snprintf 拼接时 SSID 含特殊字符导致损坏）。')
    add_code_block(doc, CODE_web_routes, '▼ API 路由注册表')
    add_code_block(doc, CODE_web_handlers, '▼ 5个代表性 API Handler 代码 (状态/配网/发送/实时流/下载)')

    # 8.9
    add_heading_styled(doc, '8.9 config_manager.c — JSON 配置管理器', level=2)
    add_normal(doc, '从 /sdcard/config.json 加载配置，缺失字段使用默认值，authmode 合法性校验，stream_gap_ms clamp 到 [1, 1000]，写盘失败自动回滚旧值。')
    add_code_block(doc, CODE_config, '▼ 配置加载代码')

    # 8.10
    add_heading_styled(doc, '8.10 console_cmd.c — 调试控制台', level=2)
    add_normal(doc, 'UART1 (TX1/RX2) 115200 baud，基于 ESP-IDF Console REPL 组件，内置 sd_test 命令。')
    add_code_block(doc, CODE_console, '▼ 调试控制台代码')

    # ============ 第9章 Web界面 ============
    add_heading_styled(doc, '9. Web 管理界面', level=1)

    add_heading_styled(doc, '9.1 访问方式', level=2)
    add_bullet(doc, '上电后自动启动 AP 热点 "ESP_Recorder_XXXX" (XXXX=MAC后两位)')
    add_bullet(doc, '密码: 12345678 (默认), 连接后浏览器访问 http://192.168.4.1')
    add_bullet(doc, '纯 HTML+CSS+Vanilla JS, 零外部依赖, 固件内嵌 (EMBED_FILES)')

    add_heading_styled(doc, '9.2 REST API 端点清单', level=2)
    add_table_with_style(doc, ['方法', 'URI', '功能说明'], [
        ['GET', '/', '返回内嵌 index.html 单页应用'],
        ['GET', '/api/status', '设备完整状态 JSON (25+字段,含UART/WiFi/SD/记录/路由)'],
        ['POST', '/api/wifi', '设置 STA WiFi 凭证 → 写入config.json → 触发重连'],
        ['POST', '/api/ap', '设置 AP SSID/密码/IP → 重启热点'],
        ['GET', '/api/files', '列出 /sdcard 下所有 REC_*.bin 文件'],
        ['GET', '/api/file?name=', '下载指定 .bin 文件 (Content-Disposition: attachment)'],
        ['POST', '/api/file/delete?name=', '删除指定文件 (路径穿越防护)'],
        ['POST', '/api/recorder/start', '开始双通道记录 (创建新REC_CH*.bin)'],
        ['POST', '/api/recorder/stop', '停止记录 (flush缓冲→fsync→fclose)'],
        ['GET', '/api/stream?since=N', '增量拉取实时流 (since_seq机制, 避免重复)'],
        ['POST', '/api/stream/clear', '清空实时流环形缓冲'],
        ['POST', '/api/send', '发送HEX字节到指定通道TX (如 01 03 00 0A)'],
        ['POST', '/api/config/stream_gap', '设置实时流间隔阈值 (1~1000ms)'],
        ['POST', '/api/reboot', '重启设备 (200ms延迟保证响应发出)'],
    ], col_widths=[2, 5.5, 8.5])

    add_heading_styled(doc, '9.3 Web 前端 JavaScript 核心代码', level=2)
    add_normal(doc, '以下展示 index.html 中 4 个最关键的 JavaScript 函数：实时流轮询、HEX发送、状态刷新、行数限制。')
    add_code_block(doc, CODE_html_js, '▼ index.html 核心 JS 代码')

    add_heading_styled(doc, '9.4 功能演示截图区', level=2)
    add_normal(doc, '请将以下功能对应的照片插入下方占位区域。', bold=True)
    add_photo_placeholder(doc, '图9-1: 连接WiFi热点 ESP_Recorder_XXXX 的手机截图')
    add_photo_placeholder(doc, '图9-2: 浏览器访问 192.168.4.1 打开的Web首页全貌')
    add_photo_placeholder(doc, '图9-3: 设备状态面板 — 显示AP已开启、SD已挂载、记录中的完整状态')
    add_photo_placeholder(doc, '图9-4: AP热点配置区 — 修改SSID/密码/IP的界面')
    add_photo_placeholder(doc, '图9-5: STA配网区 — 选择外部WiFi并输入密码')
    add_photo_placeholder(doc, '图9-6: 实时数据监控区 — HEX数据流动态滚动显示')
    add_photo_placeholder(doc, '图9-7: 文件管理区 — REC_CH1_*.bin 和 REC_CH2_*.bin 列表及下载')
    add_photo_placeholder(doc, '图9-8: 发送HEX数据功能 — 选择通道并发送的截图')
    add_photo_placeholder(doc, '图9-9: 设备控制按钮 — 开始/停止记录、重启区域')

    # ============ 第10章 数据流 ============
    add_heading_styled(doc, '10. 数据流与工作流程', level=1)

    add_heading_styled(doc, '10.1 典型使用流程', level=2)
    add_table_with_style(doc, ['阶段', '系统操作'], [
        ['1. 上电', 'ESP32-S3初始化外设→挂载SD卡→加载config.json→启动WiFi AP→(可选)连接外部WiFi→(可选)自动开始记录'],
        ['2. 开始记录', '创建 REC_CH1_YYYYMMDD_HHMMSS.bin + REC_CH2_YYYYMMDD_HHMMSS.bin → 两路UART开始接收'],
        ['3. 数据循环', 'CH1/CH2 RX → recorder环形缓冲→落盘SD卡 + router双向转发 + Web 200ms实时流推送'],
        ['4. 停止记录', '(手动/USB连接自动) 刷新剩余缓冲→fsync→fclose→通过Web下载或USB读卡器取文件'],
    ], col_widths=[2.5, 13.5])

    add_heading_styled(doc, '10.2 USB 连接时的自动行为', level=2)
    add_table_with_style(doc, ['事件', '系统响应', '用户操作'], [
        ['USB 插入电脑', 'recorder_stop()→刷新缓冲→关闭文件→sd_card_unmount_fs()卸载FATFS→SD交给USB MSC', '电脑弹出U盘,直接读写SD卡文件'],
        ['USB 拔出', 'sd_card_remount_fs()→重初始化SPI→挂载FATFS→若auto_start=1则recorder_start()恢复记录', '设备恢复串口数据记录'],
    ], col_widths=[3, 8, 5])

    # ============ 第11章 调试 ============
    add_heading_styled(doc, '11. 调试与常见问题', level=1)

    add_heading_styled(doc, '11.1 串口控制台', level=2)
    add_normal(doc, '连接 UART1 (TX1/RX2, 115200 baud)，可用命令：')
    add_code_block(doc, 'esp> sd_test                 # 标准四档测试 (1K/4K/16K/64K)\nesp> sd_test 4096 1048576      # 自定义: 4KB块, 1MB总量')

    add_heading_styled(doc, '11.2 系统日志标签', level=2)
    add_table_with_style(doc, ['日志标签', '所属模块', '典型输出'],
        [['main', '主流程', '"esp_recorder init complete"'],
         ['sd_card', 'SD卡', '"SD mounted at /sdcard, size: 7648 MB"'],
         ['usb_cdc/usb_msc', 'USB', '"TinyUSB CDC + MSC composite device ready"'],
         ['wifi_config', 'WiFi', '"AP applied: SSID=ESP_Recorder_A3F2"'],
         ['web_server', 'HTTP', '"web server started on port 80"'],
         ['recorder', '记录', '"CH1 file opened: REC_CH1_20260727_193428.bin"'],
         ['uart_driver', 'UART', '"CH1 UART0 installed @ 115200 baud"'],
         ['config', '配置', '"config loaded from /sdcard/config.json"'],
         ['data_router', '路由', '"routing enabled: CH1 RX → CH2 TX"']],
        col_widths=[3.5, 3, 9.5])

    add_heading_styled(doc, '11.3 常见问题排障', level=2)
    add_table_with_style(doc, ['故障现象', '可能原因', '排查步骤'], [
        ['Write timeout 烧录失败', 'COM口是蓝牙虚拟串口(如COM90)', '设备管理器→端口→确认ESP32的VID=303A的真实COM口'],
        ['SD卡 "not mounted"', 'SD卡未插入/接触不良/SPI接线错', '检查卡座焊接, 确认 MOSI↔MOSI, MISO↔MISO 未交叉'],
        ['WiFi AP 搜不到', 'GPIO48 VBUS检测异常或USB未接', '检查USB连接, 或确认sdkconfig中TinyUSB VBUS配置'],
        ['config.json不生效', 'JSON格式错误或路径问题', '删除后重启让系统自动生成, 或用电脑编辑验证JSON合法性'],
        ['记录文件0字节', '串口未收到数据或波特率不匹配', '检查外部设备TX→ESP32 RX接线, 确认波特率115200一致'],
        ['Web页面加载空白', '浏览器缓存或WiFi断开', '重新连接ESP_Recorder_XXXX热点, 刷新 http://192.168.4.1'],
    ], col_widths=[4, 5, 7])

    add_photo_placeholder(doc, '图11-1: 串口调试控制台 esp> sd_test 命令输出截图')
    add_photo_placeholder(doc, '图11-2: 设备管理器 COM 口列表截图 (标注 ESP32 真实串口)')

    # ============ 第12章 已知问题 ============
    add_heading_styled(doc, '12. 已知问题与注意事项', level=1)

    add_heading_styled(doc, '12.1 已验证通过', level=2)
    for s in ['ESP-IDF v5.4 环境完整配置 (esp_idf.json注册表 + VS Code扩展v2.1.0)',
              'idf.py build 编译成功 (esp32s3目标, 0 error, 0 warning)',
              '固件 esp_recorder_v1.0.bin 正常生成 (0xFA000 bytes, 76% Flash空闲)',
              'VS Code 工作区配置正确 (.vscode/settings.json 含idf.currentSetup)',
              '全部16个功能模块代码完整, 14个API端点路由注册正确']:
        add_bullet(doc, s)

    add_heading_styled(doc, '12.2 待硬件上板验证', level=2)
    for s in ['CH1 UART (TX43/RX44) 实际收发 — 需连接外部设备测试',
              'CH2 UART (TX17/RX18) 实际收发 — 需连接外部设备测试',
              'SD卡 SPI 模式实际读写 — 需插入SD卡并记录数据验证',
              'WiFi AP 手机连接 — 需上电后用手机搜索热点验证',
              'USB CDC/MSC 复合设备枚举 — 需连接电脑 USB OTG 口验证',
              '长时间记录稳定性 — 需跑几小时压力测试确认无内存泄漏',
              'USB 热插拔 FATFS/MSC 切换 — 需实际反复插拔验证可靠性']:
        add_bullet(doc, s)

    add_heading_styled(doc, '12.3 已知限制', level=2)
    add_table_with_style(doc, ['限制项', '说明', '影响程度'], [
        ['波特率固定', '当前两路共用同一波特率(115200), 不支持CH1/CH2不同速率', '低'],
        ['仅HTTP', 'Web不支持HTTPS加密, 不建议公网环境使用', '低'],
        ['Raw二进制格式', '记录文件为.bin原始字节, 非CSV, 需自行解析', '中'],
        ['无数据压缩', '长时间记录会占用较大SD空间 (115200bps≈14KB/s)', '中'],
        ['仅8N1', '串口参数固定8数据位/无校验/1停止位', '低'],
        ['无MQTT/云', '当前仅WiFi局域网访问, 无远程云平台接入', '中'],
    ], col_widths=[4, 8.5, 2])

    # ============ 附录 ============
    add_page_break(doc)
    add_heading_styled(doc, '附录 A: 源文件代码量统计', level=1)
    add_table_with_style(doc, ['文件', '大约行数', '功能'], [
        ['main/app_main.c', '183 行', '主入口, 13步系统初始化'],
        ['main/web_server.c', '649 行', 'HTTP服务器 + 14个API处理器'],
        ['main/data_recorder.c', '614 行', '双通道记录引擎 (最复杂模块)'],
        ['main/config_manager.c', '278 行', 'JSON配置读写 + 合法性校验'],
        ['main/wifi_config.c', '253 行', 'WiFi AP+STA管理 + 事件处理'],
        ['main/uart_driver.c', '221 行', '双路UART驱动 + 统计'],
        ['main/usb_msc.c', '189 行', 'USB MSC SCSI读写回调'],
        ['main/sd_card.c', '171 行', 'SD卡SPI FATFS挂载/卸载/重挂载'],
        ['main/usb_cdc.c', '150 行', 'USB CDC虚拟串口 + 描述符'],
        ['main/index.html', '309 行', 'Web前端SPA (零依赖)'],
        ['main/data_router.c', '79 行', '跨通道路由转发'],
        ['main/console_cmd.c', '91 行', '调试控制台 REPL'],
    ], col_widths=[5, 2.5, 8.5])

    add_heading_styled(doc, '附录 B: VS Code 环境恢复步骤', level=1)
    add_normal(doc, '如在新电脑继续开发，按以下步骤恢复完整开发环境：')
    add_code_block(doc, '''# 1. 确认 ESP-IDF 安装
idf.py --version          # 应输出: ESP-IDF v5.4

# 2. 确认 Python 虚拟环境
python --version          # 应输出: Python 3.8.x
python -c "import serial; print('pyserial OK')"

# 3. 确认 VS Code 扩展
code --list-extensions --show-versions | Select-String espressif
# 应输出: espressif.esp-idf-extension@2.1.0

# 4. 确认 esp_idf.json 注册表
Get-Content C:\\Users\\23178\\.espressif\\esp_idf.json

# 5. 打开项目
code D:\\zhuomian\\weite\\ESP32\\esp_recorder_v1.1

# 6. 编译验证
cd D:\\zhuomian\\weite\\ESP32\\esp_recorder_v1.1
idf.py build''')

    add_heading_styled(doc, '附录 C: 关键设计决策', level=1)
    add_table_with_style(doc, ['设计决策', '选择', '理由'], [
        ['SD卡接口', 'SPI 而非 SDMMC', '释放 SDMMC 引脚给其他外设, SPI 速率足够记录吞吐'],
        ['配置存储', 'SD卡 JSON 而非 NVS', '用户可直接用电脑编辑 config.json, 无需专用工具'],
        ['WiFi模式', 'AP+STA 混合', '保证本地 Web 管理永远可用, 同时可选外网接入'],
        ['双通道策略', '独立初始化/独立统计', '某通道硬件故障不影响另一通道正常工作'],
        ['USB方案', 'TinyUSB 复合设备', '一根USB线同时提供虚拟串口调试+读卡器取文件'],
        ['JSON拼装', 'cJSON 库', '避免 snprintf 拼接时 SSID 含引号/反斜杠等特殊字符导致损坏'],
        ['Web部署', 'EMBED_FILES 内嵌', '固件自包含, 不依赖 SD 卡上的 Web 文件'],
        ['落盘策略', '双条件触发', '环形缓冲≥2KB 或 1秒超时 → 减少写入次数又保证数据不丢'],
    ], col_widths=[3, 4.5, 8.5])

    # ============ 保存 ============
    output_path = r'D:\zhuomian\weite\ESP32\ESP_Recorder_项目开发完整报告.docx'
    doc.save(output_path)
    print(f'✅ Word 文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    build_docx()
