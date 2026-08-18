from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\zhuomian\weite\ESP32\esp_recorder_v1.1")
OUT = ROOT / "docs" / "ESP32-S3_WiFi数据记录仪_第一阶段工作验证报告.docx"
BUILD_SHOT = ROOT / "artifacts" / "01_VSCode_构建成功.png"
AUDIT_SHOT = ROOT / "artifacts" / "02_VSCode_审计报告.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E0"
GOLD = "7A5A00"
PALE_RED = "FDECEC"
RED = "9B1C1C"
GREEN = "166534"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def set_run_font(run, size=None, bold=None, color=INK, italic=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, before=0, after=6, line=1.1, align=None, keep_next=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep_next is not None:
        pf.keep_with_next = keep_next


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        set_row_cant_split(row)
        for idx, width in enumerate(widths_dxa):
            set_cell_width(row.cells[idx], width)
            set_cell_margins(row.cells[idx])


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_cell_text(cell, size=9.2, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        set_para(p, before=0, after=0, line=1.12, align=align)
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None, color=INK, after=6):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.1)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    set_para(p, after=4, line=1.167)
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    set_para(p, after=4, line=1.167)
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    set_para(p, before=0, after=0, line=1.1)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, bold=True, color=label_color)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def add_picture_with_alt(doc, path, width_inches, title, description):
    p = doc.add_paragraph()
    set_para(p, before=4, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=0, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=False)
    r = p.add_run(text)
    set_run_font(r, size=9.2, color=MUTED)
    return p


def add_manual_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[list_style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.167

header = section.header
hp = header.paragraphs[0]
set_para(hp, after=0, line=1.0)
left = hp.add_run("ESP32-S3 WiFi 数据记录仪")
set_run_font(left, size=9, bold=True, color=MUTED)
right = hp.add_run("  |  第一阶段验证报告")
set_run_font(right, size=9, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)

# Cover / memo masthead.
p = doc.add_paragraph()
set_para(p, before=18, after=4, line=1.0)
r = p.add_run("技术验证报告")
set_run_font(r, size=11, bold=True, color=BLUE)

p = doc.add_paragraph()
set_para(p, before=0, after=6, line=1.0)
r = p.add_run("ESP32-S3 WiFi 数据记录仪")
set_run_font(r, size=25, bold=True, color=INK)

p = doc.add_paragraph()
set_para(p, before=0, after=20, line=1.0)
r = p.add_run("第一阶段代码接手、构建验证与功能证据记录")
set_run_font(r, size=14, color=MUTED)

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(meta, [1600, 7760])
set_table_borders(meta, color=WHITE, size=0)
meta_rows = [
    ("项目目录", r"D:\zhuomian\weite\ESP32\esp_recorder_v1.1"),
    ("项目版本", "esp_recorder_v1.1（工程名仍为 esp_recorder_v1.0）"),
    ("目标平台", "ESP32-S3 / ESP-IDF v5.4"),
    ("验证日期", "2026-07-27"),
    ("报告状态", "第一阶段完成：代码审计与实际构建；未执行烧录"),
]
for i, (label, value) in enumerate(meta_rows):
    meta.cell(i, 0).text = label
    meta.cell(i, 1).text = value
    set_cell_shading(meta.cell(i, 0), LIGHT)
    style_cell_text(meta.cell(i, 0), size=9.8, color=DARK_BLUE, bold=True)
    style_cell_text(meta.cell(i, 1), size=9.8, color=INK)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_callout(
    doc,
    "结论",
    "当前代码在 ESP-IDF v5.4、目标 esp32s3 配置下构建成功并生成固件；这仅证明工程可构建，WiFi、UART、SD 和 USB 行为仍需开发板实测。",
    fill=PALE_BLUE,
)
add_callout(
    doc,
    "边界",
    "遵守阶段要求：未修改业务源码，未执行 set-target、erase-flash 或 flash。VS Code 配置的 COM25 当前不存在，不能把本轮结果描述为已上板运行通过。",
    fill=PALE_GOLD,
    label_color=GOLD,
)

add_manual_page_break(doc)

add_heading(doc, "1. 工作目标与执行范围", 1)
add_body(
    doc,
    "本报告用于记录现有 ESP32-S3 WiFi 数据记录仪项目的接手审计、开发环境核对、真实构建、七项功能调用链检查、风险识别和成功证据归档。"
)
add_body(doc, "本轮执行内容包括：", bold_prefix="本轮执行内容包括：")
for item in (
    "核对项目目录、Git 状态、ESP-IDF 环境、目标芯片和 VS Code 配置。",
    "完整读取项目自有源码、配置文件、网页资源及 TinyUSB 集成入口。",
    "使用现有 ESP-IDF v5.4 环境实际执行 IDF build。",
    "逐项检查 WiFi、网页、UART、SD 记录、波特率和 USB MSC 的真实调用链。",
    "保存完整构建日志、VS Code 成功截图和阶段审计报告。",
):
    add_bullet(doc, item)

add_heading(doc, "2. 项目与环境检查", 1)
env_table = doc.add_table(rows=1, cols=3)
env_table.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ["检查项", "实际结果", "判定"]
for i, text in enumerate(headers):
    env_table.cell(0, i).text = text
    set_cell_shading(env_table.cell(0, i), LIGHT)
    style_cell_text(env_table.cell(0, i), size=9.5, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
env_rows = [
    ("Git", "当前目录不是 Git 仓库", "存在回退风险"),
    ("ESP-IDF", "ESP-IDF v5.4；组件版本 5.4.0", "可用于构建"),
    ("目标芯片", 'sdkconfig: CONFIG_IDF_TARGET="esp32s3"', "符合预期"),
    ("VS Code 工具链", "配置路径指向不存在的 C:/E: 目录；实际环境在 D:/23178 与用户 .espressif", "配置需修正"),
    ("串口", "配置 COM25；当前仅枚举 COM4/6/7/90", "未识别开发板端口"),
    ("分区", "实际使用 1 MB App 分区；项目 partitions.csv 未启用", "配置矛盾"),
]
for row_data in env_rows:
    cells = env_table.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].text = text
        style_cell_text(cells[i], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT)
set_repeat_table_header(env_table.rows[0])
set_table_geometry(env_table, [1750, 5550, 2060])
set_table_borders(env_table)

add_manual_page_break(doc)

add_heading(doc, "3. 实际构建验证", 1)
add_body(doc, "构建命令：", bold_prefix="构建命令：")
p = doc.add_paragraph()
set_para(p, before=0, after=8, line=1.0)
r = p.add_run(
    r"C:\Users\23178\.espressif\python_env\idf5.4_py3.8_env\Scripts\python.exe "
    r"D:\23178\esp-idf\tools\idf.py build"
)
set_run_font(r, size=8.8, color=DARK_BLUE, font="Consolas")

build_table = doc.add_table(rows=1, cols=2)
build_table.alignment = WD_TABLE_ALIGNMENT.LEFT
build_rows = [
    ("构建结果", "成功，完成 1066 个构建步骤"),
    ("固件文件", r"build\esp_recorder_v1.0.bin"),
    ("固件大小", "0xfbb40 字节（1,030,976 字节）"),
    ("分区余量", "0x44c0 字节，约 2%"),
    ("固件 SHA-256", "904F3F7A84195F87854A5AB8752C0DBADC992D23623DB6B54EDAAD50E4EA5B0B"),
    ("完整日志", r"artifacts\idf_build_20260727.txt"),
]
build_table.cell(0, 0).text = build_rows[0][0]
build_table.cell(0, 1).text = build_rows[0][1]
for label, value in build_rows[1:]:
    cells = build_table.add_row().cells
    cells[0].text = label
    cells[1].text = value
for row in build_table.rows:
    set_cell_shading(row.cells[0], LIGHT)
    style_cell_text(row.cells[0], size=9.4, color=DARK_BLUE, bold=True)
    style_cell_text(row.cells[1], size=9.1)
set_table_geometry(build_table, [1900, 7460])
set_table_borders(build_table)

add_callout(
    doc,
    "重要警告",
    "构建工具明确提示最小 App 分区接近满载，仅剩约 2%。后续很小的代码增长也可能导致链接失败，应先确认 Flash 容量并统一 sdkconfig 与 partitions.csv。",
    fill=PALE_RED,
    label_color=RED,
)

add_manual_page_break(doc)

add_heading(doc, "4. 七项功能验证结论", 1)
add_body(
    doc,
    "状态说明：B 表示当前配置下编译通过但仍需硬件实测；C 表示只有部分链路实现；F 表示存在可由代码确认的明显缺陷。"
)
matrix = doc.add_table(rows=1, cols=4)
matrix.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, text in enumerate(("序号", "功能", "状态", "审计结论")):
    matrix.cell(0, i).text = text
    set_cell_shading(matrix.cell(0, i), PALE_BLUE)
    style_cell_text(matrix.cell(0, i), size=9.2, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
matrix_rows = [
    ("1", "WiFi AP + Station 混合模式", "B", "AP/STA netif、APSTA 模式、WiFi/IP 事件和重连链路均存在并完成构建；掉线回调内 5 秒延迟会阻塞默认事件循环。"),
    ("2", "USB 或 AP 网页配网", "C", "后端真实接口为 POST /api/wifi，但嵌入前端调用 /api/wifi/connect、disconnect、scan；USB CDC 回调直接丢弃数据。"),
    ("3", "手机连接 AP 后访问网页", "B", "HTTP Server、根路由和嵌入式 HTML/CSS/JS 均参与构建；需开发板验证 192.168.4.1、静态资源和异常响应。"),
    ("4", "单串口实时监控与双向透传", "C", "RX 与网页发送链路存在；TX 未接入 live，USB CDC 不转发 UART，且用户串口与 ESP Console 都占用 UART0。"),
    ("5", "局域网开始/停止 SD 记录", "F", "记录链路可构建，但 recorder_stop 先清除运行状态，最终 ring buffer flush 会被跳过，停止时可能丢失尾部数据。"),
    ("6", "串口波特率配置", "C", "启动时能读取 baudrate；网页提交的 /api/uart/config 后端未注册，运行时设置与持久化没有闭环。"),
    ("7", "Type-C 模拟 U 盘", "B", "TinyUSB MSC 回调直接读写 SD 扇区并构建成功；自定义描述符、GPIO48 VBUS 和 FATFS/MSC 互斥仍需实机确认。"),
]
status_fill = {"B": PALE_BLUE, "C": PALE_GOLD, "F": PALE_RED}
status_color = {"B": DARK_BLUE, "C": GOLD, "F": RED}
for row_data in matrix_rows:
    cells = matrix.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].text = text
        style_cell_text(cells[i], size=8.65, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_shading(cells[2], status_fill[row_data[2]])
    style_cell_text(cells[2], size=9.2, color=status_color[row_data[2]], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
set_repeat_table_header(matrix.rows[0])
set_table_geometry(matrix, [620, 2200, 700, 5840])
set_table_borders(matrix)

add_manual_page_break(doc)

add_heading(doc, "5. 主要风险与问题记录", 1)
risks = [
    ("P1", "UART0 资源冲突", "用户串口定义为 UART0，sdkconfig 同时把 ESP Console 配置为 UART0；REPL 初始化可能失败并触发 ESP_ERROR_CHECK。"),
    ("P1", "停止记录丢数据", "recorder_stop 的状态顺序导致最终环形缓冲刷新被跳过，停止操作不能保证尾部数据落盘。"),
    ("P1", "前后端接口不一致", "网页 WiFi 与 UART 按钮调用未注册路由，页面外观不能证明配置功能闭环。"),
    ("P2", "分区与 Flash 配置矛盾", "项目 partitions.csv 未启用，当前 1 MB App 分区仅剩 2%，后续维护空间不足。"),
    ("P2", "USB 描述符与硬件依赖", "Kconfig 启用 CDC+MSC，但自定义 configuration descriptor 只声明 MSC；GPIO48 VBUS 连接也缺少原理图证据。"),
]
risk_table = doc.add_table(rows=1, cols=3)
for i, text in enumerate(("级别", "问题", "影响")):
    risk_table.cell(0, i).text = text
    set_cell_shading(risk_table.cell(0, i), LIGHT)
    style_cell_text(risk_table.cell(0, i), size=9.3, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for level, title, impact in risks:
    cells = risk_table.add_row().cells
    cells[0].text = level
    cells[1].text = title
    cells[2].text = impact
    style_cell_text(cells[0], size=9.2, color=RED if level == "P1" else GOLD, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell_text(cells[1], size=9.2, bold=True)
    style_cell_text(cells[2], size=8.9)
set_repeat_table_header(risk_table.rows[0])
set_table_geometry(risk_table, [820, 2280, 6260])
set_table_borders(risk_table)

add_manual_page_break(doc)

add_heading(doc, "6. 成功证据截图", 1)
add_body(
    doc,
    "截图已逐张复核。以下两张图片能对应本轮真实 VS Code 状态和构建结果，因此纳入报告。原有 artifacts/dashboard.png 显示“127.0.0.1 拒绝连接”，属于失败页面，已确认不作为成功证据，也未插入本报告。"
)
add_heading(doc, "6.1 VS Code 构建成功证据", 2)
add_body(
    doc,
    "截图中可见 ESP_RECORDER_V1.1 工程树、idf_build_20260727.txt、固件大小与分区余量警告，以及“Project build complete”结果。"
)
add_picture_with_alt(
    doc,
    BUILD_SHOT,
    6.15,
    "VS Code 构建成功",
    "VS Code 打开 ESP32 工程和构建日志，日志显示 Project build complete，并提示 App 分区仅剩 2%。",
)
add_caption(doc, "图 1  VS Code 中的 IDF 构建成功日志（有效成功证据）")

add_manual_page_break(doc)

add_heading(doc, "6.2 VS Code 审计报告证据", 2)
add_body(
    doc,
    "截图中可见 VS Code 工程树、构建日志标签和 HANDOVER_AUDIT_STAGE1.md，报告正文显示目标、构建结果和七项功能验证矩阵。"
)
add_picture_with_alt(
    doc,
    AUDIT_SHOT,
    5.65,
    "VS Code 审计报告",
    "VS Code 打开第一阶段审计报告，显示 ESP32-S3 目标、构建结果和七项功能验证矩阵。",
)
add_caption(doc, "图 2  VS Code 中的第一阶段审计报告（有效工作记录证据）")

add_heading(doc, "7. 硬件实测边界", 1)
add_callout(
    doc,
    "不能误报",
    "本轮没有烧录开发板，且 COM25 不存在，因此不能宣称 AP、STA、UART、SD 卡或 USB U 盘已在实体硬件上运行成功。当前确认的是源码调用链和构建结果。",
    fill=PALE_GOLD,
    label_color=GOLD,
)
add_body(doc, "后续硬件验证应按以下顺序执行：")
for item in (
    "确认开发板型号、Flash/PSRAM 容量、原理图和实际串口号。",
    "修复 UART0/Console 冲突并明确用户串口与调试口。",
    "修复 recorder_stop 最终刷新顺序，执行短包、满速、停止和断电测试。",
    "对齐网页与后端 WiFi/UART API，再验证 AP 页面和 STA 重连。",
    "确认分区表与 Flash 容量，给 App 保留合理升级余量。",
    "最后验证 USB MSC 枚举、读写、弹出、断连重挂载与记录互斥。",
):
    add_number(doc, item)

add_heading(doc, "8. 本轮交付物", 1)
deliverables = [
    ("完整构建日志", r"artifacts\idf_build_20260727.txt"),
    ("VS Code 构建成功截图", r"artifacts\01_VSCode_构建成功.png"),
    ("VS Code 审计报告截图", r"artifacts\02_VSCode_审计报告.png"),
    ("Markdown 审计报告", r"docs\HANDOVER_AUDIT_STAGE1.md"),
    ("Word 工作验证报告", r"docs\ESP32-S3_WiFi数据记录仪_第一阶段工作验证报告.docx"),
]
deliv_table = doc.add_table(rows=1, cols=2)
deliv_table.cell(0, 0).text = "交付物"
deliv_table.cell(0, 1).text = "相对路径"
for c in deliv_table.rows[0].cells:
    set_cell_shading(c, LIGHT)
    style_cell_text(c, size=9.3, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for name, path in deliverables:
    cells = deliv_table.add_row().cells
    cells[0].text = name
    cells[1].text = path
    style_cell_text(cells[0], size=9.2, bold=True)
    style_cell_text(cells[1], size=8.9)
set_table_geometry(deliv_table, [2900, 6460])
set_table_borders(deliv_table)

add_heading(doc, "附录 A：构建结束关键信息", 1)
lines = [
    "Generated D:/zhuomian/weite/ESP32/esp_recorder_v1.1/build/esp_recorder_v1.0.bin",
    "[1066/1066] ... check_sizes.py ... esp_recorder_v1.0.bin",
    "esp_recorder_v1.0.bin binary size 0xfbb40 bytes.",
    "Smallest app partition is 0x100000 bytes.",
    "0x44c0 bytes (2%) free.",
    "Warning: The smallest app partition is nearly full (2% free space left)!",
    "Project build complete. To flash, run:",
    " idf.py flash",
]
code_table = doc.add_table(rows=1, cols=1)
set_table_geometry(code_table, [9360])
set_table_borders(code_table, color=BORDER, size=4)
cell = code_table.cell(0, 0)
set_cell_shading(cell, "F8FAFC")
set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
cell.paragraphs[0].clear()
for idx, line in enumerate(lines):
    p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
    set_para(p, before=0, after=1, line=1.0)
    r = p.add_run(line)
    set_run_font(r, size=8.3, color=INK, font="Consolas")

add_callout(
    doc,
    "说明",
    "“To flash”是构建工具自动给出的后续提示。本轮没有执行 idf.py flash、erase-flash 或任何等价烧录命令。",
    fill=PALE_BLUE,
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
