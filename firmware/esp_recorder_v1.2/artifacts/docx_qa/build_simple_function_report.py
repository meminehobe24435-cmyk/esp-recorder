from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\zhuomian\weite\ESP32\esp_recorder_v1.1")
OUT = ROOT / "docs" / "ESP32-S3功能验证报告_简版.docx"
BUILD_SHOT = ROOT / "artifacts" / "05_VSCode_重新构建成功.png"
MSC_SHOT = ROOT / "artifacts" / "04_USB_MSC_U盘读取成功_报告裁剪.png"

BLUE = "2E74B5"
DARK = "1F2937"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF6EE"
PALE_GOLD = "FFF7E0"
GREEN = "166534"
GOLD = "7A5A00"
BORDER = "CBD5E1"
WHITE = "FFFFFF"


def font(run, size=10.5, bold=False, color=DARK, name="Microsoft YaHei"):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(p, before=0, after=5, line=1.15, align=None, keep_next=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep_next is not None:
        p.paragraph_format.keep_with_next = keep_next


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def borders(table):
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.find(qn("w:tblBorders"))
    if node is None:
        node = OxmlElement("w:tblBorders")
        tbl_pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "5")
        item.set(qn("w:color"), BORDER)
        node.append(item)


def set_cell(cell, text, size=9.2, bold=False, color=DARK, align=None):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        para(p, after=0, line=1.05, align=align)
        for run in p.runs:
            font(run, size=size, bold=bold, color=color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    font(r, size=16 if level == 1 else 12, bold=True, color=BLUE)
    para(p, before=8 if level == 1 else 4, after=6, keep_next=True)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        font(r2)
    else:
        r = p.add_run(text)
        font(r)
    para(p)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    font(r, size=10)
    para(p, after=3)


def picture(doc, path, width, caption, alt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", alt)
    para(p, after=3)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption)
    font(r, size=9, color=MUTED)
    para(cp, after=6)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)

for sec in doc.sections:
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("ESP32-S3 WiFi 数据记录仪  |  功能验证")
    font(run, size=8.5, bold=True, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("验证日期：2026-07-27")
    font(run, size=8, color=MUTED)

# Cover
p = doc.add_paragraph()
r = p.add_run("功能验证报告")
font(r, size=15, bold=True, color=BLUE)
para(p, before=28, after=10)

p = doc.add_paragraph()
r = p.add_run("ESP32-S3 WiFi 数据记录仪")
font(r, size=27, bold=True)
para(p, after=7)

p = doc.add_paragraph()
r = p.add_run("简版 · 构建与已连接硬件验证")
font(r, size=14, color=MUTED)
para(p, after=24)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
items = [
    ("项目目录", r"D:\zhuomian\weite\ESP32\esp_recorder_v1.1"),
    ("开发环境", "ESP-IDF v5.4 / ESP32-S3"),
    ("硬件连接", "Type-C USB MSC + CH340 USB-TTL（COM12）"),
    ("验证原则", "不烧录、不修改业务代码，只记录实际通过项"),
]
for row, (key, value) in zip(meta.rows, items):
    set_cell(row.cells[0], key, bold=True, color=BLUE)
    shade(row.cells[0], LIGHT)
    set_cell(row.cells[1], value)
    cant_split(row)
borders(meta)

doc.add_paragraph()
callout = doc.add_table(rows=1, cols=1)
set_cell(
    callout.cell(0, 0),
    "结论：工程重新构建成功；ESP32-S3 原生 USB U 盘枚举成功并可读取 SD 卡文件；"
    "CH340 COM12 可打开。UART 10 秒未收到数据，WiFi AP 未检测到，因此不判定相关数据功能通过。",
    size=10.5,
    bold=True,
)
shade(callout.cell(0, 0), PALE_BLUE)
borders(callout)

heading(doc, "1. 验证步骤", 1)
for text in (
    "加载 ESP-IDF v5.4 环境，确认 sdkconfig 目标为 esp32s3。",
    "执行 idf.py build，检查最终固件生成结果。",
    "检查原生 USB 设备 VID_303A / PID_4002，并读取 F: 文件。",
    "检查 CH340 USB-TTL 的 COM12，按 460800、8N1 打开串口 10 秒。",
):
    bullet(doc, text)

heading(doc, "2. 功能验证结果", 1)
table = doc.add_table(rows=1, cols=4)
headers = ("项目", "状态", "实际证据", "结论")
for cell, text in zip(table.rows[0].cells, headers):
    set_cell(cell, text, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(cell, LIGHT)
repeat_header(table.rows[0])

rows = [
    ("工程构建", "通过", "1066/1066；Project build complete", "固件生成成功"),
    ("Type-C U 盘", "通过", "VID_303A/PID_4002；F: FAT32", "MSC 枚举成功"),
    ("SD 文件读取", "通过", "config.json、REC_*.bin、测速文件可读", "读取成功"),
    ("USB-TTL 连接", "通过", "CH340 COM12；460800、8N1 可打开", "连接层成功"),
    ("UART 实时数据", "未通过", "10 秒接收 0 字节", "不能认定数据链路成功"),
    ("WiFi/AP 网页", "未验证", "电脑未检测到设备 AP", "需继续实机测试"),
    ("SD 开始/停止记录", "未验证", "仅确认已有文件可读", "未执行记录闭环"),
]
for item, status, evidence, conclusion in rows:
    cells = table.add_row().cells
    for i, text in enumerate((item, status, evidence, conclusion)):
        color = GREEN if status == "通过" and i == 1 else (GOLD if i == 1 else DARK)
        set_cell(cells[i], text, size=8.8, bold=(i in (0, 1)), color=color)
    if status == "通过":
        shade(cells[1], PALE_GREEN)
    else:
        shade(cells[1], PALE_GOLD)
    cant_split(table.rows[-1])
borders(table)

doc.add_page_break()
heading(doc, "3. 构建成功截图", 1)
body(doc, "步骤：在项目目录执行 idf.py build。结果：ESP32-S3 固件生成成功。")
picture(
    doc,
    BUILD_SHOT,
    6.7,
    "图 1  VS Code 中的重新构建成功结果",
    "VS Code 显示完整构建日志末尾，包含 Project build complete。",
)

doc.add_page_break()
heading(doc, "4. USB U 盘读取成功截图", 1)
body(doc, "步骤：保持 Type-C 连接，在资源管理器打开 F:。结果：FAT32 文件系统和记录文件可读取。")
picture(
    doc,
    MSC_SHOT,
    6.7,
    "图 2  ESP Recorder USB MSC 映射为 F: 并成功读取文件",
    "Windows 文件资源管理器打开 F:，显示 config.json、REC 记录文件和 SD 测速文件。",
)

heading(doc, "5. 本轮结论", 1)
body(
    doc,
    "本轮实际通过：工程构建、USB MSC 枚举、SD 文件读取、USB-TTL 端口连接。"
    "未获得成功证据的 WiFi、UART 数据收发和 SD 记录控制未标记为通过。",
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
